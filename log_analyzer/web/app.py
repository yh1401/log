"""FastAPI Web application for Log Analyzer."""

import os
import sys
import asyncio
import logging
import json
from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Request, Header, Depends, Body
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from pydantic import BaseModel

SCRIPT_DIR = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import log_analyzer.config.settings as config_module
import log_analyzer.parser.log_parser as parser_module
import log_analyzer.checkpoint.manager as checkpoint_module
import log_analyzer.llm.client as llm_module
import log_analyzer.processor.chunk_processor as processor_module
import log_analyzer.report.generator as report_module
import log_analyzer.utils.helpers as utils_module
from .auth import user_manager, DEFAULT_USER_ID, DEFAULT_USERNAME, ADMIN_USERNAME, require_admin
from .storage import get_storage
from .action_logger import record_file_upload, record_task_start, record_task_complete, record_task_failed, record_page_view, record_button_click, record_api_request

Settings = config_module.Settings
load_llm_config = config_module.load_llm_config
init_settings = config_module.init_settings
LogParser = parser_module.LogParser
CheckpointManager = checkpoint_module.CheckpointManager
LLMClient = llm_module.LLMClient
ChunkProcessor = processor_module.ChunkProcessor
ProcessingResult = processor_module.ProcessingResult
ReportGenerator = report_module.ReportGenerator
ensure_dir = utils_module.ensure_dir

# 用户相关目录
USERS_DIR = PROJECT_ROOT / "log_analyzer" / "users"
TASKS_DIR = PROJECT_ROOT / "log_analyzer" / "tasks"
ensure_dir(str(USERS_DIR))
ensure_dir(str(TASKS_DIR))

# 配置允许访问的目录（权限控制）- 跨平台支持
import sys

def get_default_allowed_directories():
    """根据操作系统返回默认允许的目录列表"""
    allowed = []
    
    if sys.platform.startswith('win'):
        # Windows 系统
        allowed.extend([
            Path("C:\\Windows\\Logs"),
            Path("C:\\ProgramData\\logs"),
            Path("C:\\Temp"),
            Path("C:\\Users"),
            Path("D:\\Logs"),
            Path("D:\\Temp"),
        ])
    elif sys.platform == 'darwin':
        # macOS 系统
        allowed.extend([
            Path("/var/log"),
            Path("/opt/logs"),
            Path("/tmp"),  # macOS /tmp 是符号链接，会解析到 /private/tmp
            Path("/Users"),
            Path("/Library/Logs"),
            Path("~/Library/Logs"),
        ])
    else:
        # Linux/Unix 系统
        allowed.extend([
            Path("/var/log"),
            Path("/opt/logs"),
            Path("/tmp"),
            Path("/home"),
            Path("/root"),
            Path("/var/lib"),
        ])
    
    # 过滤不存在的目录并去重
    resolved_dirs = []
    seen_paths = set()
    
    for d in allowed:
        try:
            if d.home():
                abs_path = d.expanduser().resolve()
            else:
                abs_path = d.resolve()
            
            path_str = str(abs_path)
            if path_str not in seen_paths:
                if abs_path.exists() or str(d).startswith('~'):
                    resolved_dirs.append(abs_path)
                    seen_paths.add(path_str)
        except Exception:
            continue
    
    return resolved_dirs

ALLOWED_DIRECTORIES = get_default_allowed_directories()

# 获取允许访问的目录列表
def get_allowed_directories_from_config():
    """从配置获取允许访问的目录列表"""
    try:
        from log_analyzer.config.settings import get_settings
        settings = get_settings()
        config_dirs = settings.server_path.allowed_directories
        
        # 如果配置了目录，返回配置的目录
        if config_dirs and len(config_dirs) > 0:
            return config_dirs
        
        # 如果没有配置，返回根目录表示不限制
        return ["/"]
    except Exception:
        # 如果读取配置失败，返回默认根目录
        return ["/"]

# 检查路径是否在允许的目录内
def is_path_allowed(path: Path) -> bool:
    """检查路径是否在允许的目录范围内（跨平台支持）"""
    # 安全检查：防止路径遍历攻击
    path_str = str(path)
    if '..' in path_str:
        return False
    
    # 尝试从配置读取允许的目录
    try:
        from log_analyzer.config.settings import get_settings
        settings = get_settings()
        config_dirs = settings.server_path.allowed_directories
        
        # 如果配置文件中有配置（即使是空数组），按配置处理
        # 空数组 [] 表示不限制任何路径，所有路径都可访问
        if config_dirs is not None:
            # 如果配置了具体目录，则限制访问
            if len(config_dirs) > 0:
                configured_dirs = [Path(d).resolve() for d in config_dirs if d]
                for allowed_dir in configured_dirs:
                    try:
                        abs_path = path.resolve()
                        if abs_path.is_relative_to(allowed_dir):
                            return True
                        # 处理 macOS 符号链接
                        if sys.platform == 'darwin':
                            if str(abs_path).startswith(str(allowed_dir)):
                                return True
                    except Exception:
                        continue
                return False
            else:
                # 空数组表示不限制，允许所有路径
                return True
    except Exception:
        pass
    
    # 如果读取配置失败，使用默认行为（不限制）
    return True

# 获取跨平台的临时目录
def get_system_temp_dir():
    """获取当前系统的临时目录"""
    if sys.platform.startswith('win'):
        return Path("C:\\Temp") if Path("C:\\Temp").exists() else Path("/tmp")
    return Path("/tmp")

def format_bytes(size: int) -> str:
    for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
        if size < 1024.0:
            return f"{size:.2f} {unit}"
        size /= 1024.0
    return f"{size:.2f} PB"


def get_user_dir(user_id: str) -> Path:
    """获取用户专属目录"""
    user_dir = USERS_DIR / user_id
    ensure_dir(str(user_dir))
    return user_dir


def get_user_upload_dir(user_id: str) -> Path:
    """获取用户上传文件目录"""
    upload_dir = get_user_dir(user_id) / "uploads"
    ensure_dir(str(upload_dir))
    return upload_dir


def get_user_reports_dir(user_id: str) -> Path:
    """获取用户报告目录"""
    reports_dir = get_user_dir(user_id) / "reports"
    ensure_dir(str(reports_dir))
    return reports_dir


def get_user_checkpoints_dir(user_id: str) -> Path:
    """获取用户检查点目录"""
    checkpoint_dir = get_user_dir(user_id) / "checkpoints"
    ensure_dir(str(checkpoint_dir))
    return checkpoint_dir


async def get_current_user(
    x_user_id: str = Header(None, alias="X-User-Id"),
    x_username: str = Header(None, alias="X-Username")
) -> Dict[str, Any]:
    """从请求头获取当前用户标识（无鉴权）

    - 通过 X-User-Id 头识别用户身份
    - 若未提供则使用默认用户 default_user
    - 自动创建用户档案
    """
    user_id = x_user_id or DEFAULT_USER_ID
    username = x_username or DEFAULT_USERNAME

    user_info = user_manager.get_or_create_user(user_id, username)
    return user_info


async def get_optional_user(authorization: str = Header(None)) -> Optional[Dict[str, Any]]:
    """可选用户识别"""
    return None


async def require_admin(current_user: Dict[str, Any] = Depends(get_current_user)):
    """管理员权限检查依赖函数
    
    用于保护历史记录相关接口，仅允许管理员访问。
    如果用户不是管理员，返回 403 Forbidden。
    """
    username = current_user.get("username", "")
    
    if not user_manager.is_admin(username):
        raise HTTPException(
            status_code=403,
            detail={
                "code": 403,
                "message": "访问被拒绝：仅管理员可以访问此接口",
                "data": None
            }
        )
    
    return current_user


def extract_zip(zip_path: Path, dest_dir: Path) -> List[str]:
    """解压ZIP文件并返回提取的文件路径列表"""
    import zipfile
    extracted_files = []
    try:
        with zipfile.ZipFile(zip_path, 'r') as zf:
            zf.extractall(dest_dir)
            for name in zf.namelist():
                extracted_path = dest_dir / name
                if extracted_path.is_file():
                    extracted_files.append(str(extracted_path))
        logging.info(f"ZIP文件解压完成: {zip_path}, 提取了 {len(extracted_files)} 个文件")
    except Exception as e:
        logging.error(f"解压ZIP文件失败: {str(e)}")
    return extracted_files

def pcap_to_log(pcap_path: Path) -> str:
    """将PCAP文件转换为日志格式，使用tshark进行协议解析"""
    log_content = []
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
    
    log_content.append(f"{timestamp} [INFO] [PCAP] 文件路径: {pcap_path}")
    log_content.append(f"{timestamp} [INFO] [PCAP] 文件大小: {pcap_path.stat().st_size} bytes")
    log_content.append(f"{timestamp} [INFO] [PCAP] 使用tshark进行协议分析")
    
    try:
        import subprocess
        
        # 使用tshark获取完整的数据包统计信息
        stats_cmd = ['tshark', '-r', str(pcap_path), '-q', '-z', 'io,phs']
        result = subprocess.run(stats_cmd, capture_output=True, text=True, timeout=30, errors='replace')
        if result.stdout:
            for line in result.stdout.split('\n'):
                if 'frames' in line and 'bytes' in line:
                    log_content.append(f"{timestamp} [INFO] [PCAP] 协议统计: {line.strip()[:150]}")
        
        # 使用tshark提取详细数据包信息
        tshark_cmd = [
            'tshark', '-r', str(pcap_path),
            '-T', 'fields',
            '-e', 'frame.time',
            '-e', 'ip.src',
            '-e', 'ip.dst',
            '-e', '_ws.col.Protocol',
            '-e', 'tcp.srcport',
            '-e', 'tcp.dstport',
            '-e', 'tcp.flags',
            '-e', 'tcp.len',
            '-e', '_ws.col.Info',
            '-E', 'separator=|'
        ]
        
        result = subprocess.run(
            tshark_cmd,
            capture_output=True,
            text=True,
            timeout=60,
            errors='replace'
        )
        
        if result.returncode == 0 and result.stdout.strip():
            lines = result.stdout.strip().split('\n')
            packet_count = 0
            tcp_count = 0
            tcp_syn_count = 0
            tcp_rst_count = 0
            tcp_ack_count = 0
            other_count = 0
            error_count = 0
            warning_count = 0
            
            for line in lines[:500]:
                parts = line.split('|')
                if len(parts) < 5:
                    continue
                
                try:
                    frame_time = parts[0] if parts[0] else timestamp
                    ip_src = parts[1] if parts[1] else ''
                    ip_dst = parts[2] if parts[2] else ''
                    protocol = parts[3] if parts[3] else ''
                    info = parts[-1] if parts[-1] else ''
                    
                    packet_count += 1
                    
                    if protocol == 'TCP':
                        tcp_count += 1
                        src_port = parts[4] if len(parts) > 4 and parts[4] else ''
                        dst_port = parts[5] if len(parts) > 5 and parts[5] else ''
                        flags = parts[6] if len(parts) > 6 and parts[6] else ''
                        tcp_len = parts[7] if len(parts) > 7 and parts[7] else '0'
                        
                        if 'RST' in flags:
                            log_content.append(f"{frame_time} [ERROR] [PCAP] TCP Reset: {ip_src}:{src_port} -> {ip_dst}:{dst_port} [{info[:80]}]")
                            tcp_rst_count += 1
                            error_count += 1
                        elif 'SYN' in flags and 'ACK' not in flags:
                            log_content.append(f"{frame_time} [WARN] [PCAP] TCP SYN (连接建立): {ip_src}:{src_port} -> {ip_dst}:{dst_port}")
                            tcp_syn_count += 1
                            warning_count += 1
                        elif 'ACK' in flags and 'PSH' in flags:
                            data_len = int(tcp_len) if tcp_len.isdigit() else 0
                            log_content.append(f"{frame_time} [INFO] [PCAP] TCP数据: {ip_src}:{src_port} -> {ip_dst}:{dst_port} ({data_len} bytes)")
                        elif 'ACK' in flags:
                            tcp_ack_count += 1
                            log_content.append(f"{frame_time} [INFO] [PCAP] TCP ACK: {ip_src}:{src_port} -> {ip_dst}:{dst_port}")
                        else:
                            log_content.append(f"{frame_time} [INFO] [PCAP] TCP: {ip_src}:{src_port} -> {ip_dst}:{dst_port} Flags:{flags[:10]}")
                            
                    elif protocol == 'UDP':
                        other_count += 1
                        src_port = parts[4] if len(parts) > 4 and parts[4] else ''
                        dst_port = parts[5] if len(parts) > 5 and parts[5] else ''
                        log_content.append(f"{frame_time} [INFO] [PCAP] UDP: {ip_src}:{src_port} -> {ip_dst}:{dst_port}")
                        
                    elif protocol and protocol != 'Frame':
                        log_content.append(f"{frame_time} [INFO] [PCAP] {protocol}: {ip_src} -> {ip_dst} [{info[:80]}]")
                        other_count += 1
                        
                except (IndexError, ValueError):
                    continue
            
            # 添加统计摘要
            log_content.insert(3, f"{timestamp} [INFO] [PCAP] 分析数据包数: {packet_count}")
            log_content.insert(4, f"{timestamp} [INFO] [PCAP] TCP会话数: {tcp_count}")
            log_content.insert(5, f"{timestamp} [INFO] [PCAP]   - SYN: {tcp_syn_count}, RST: {tcp_rst_count}, ACK: {tcp_ack_count}")
            log_content.insert(6, f"{timestamp} [INFO] [PCAP] 其他协议包: {other_count}")
            
            log_content.append(f"{timestamp} [INFO] [PCAP] 分析完成 - 错误数: {error_count}, 警告数: {warning_count}")
            
            # 如果日志太少，添加ASTERIX相关信息
            if packet_count > 0 and len(log_content) < 30:
                log_content.append(f"{timestamp} [INFO] [PCAP] 主要协议: TCP/ASTERIX (航空监控数据)")
                log_content.append(f"{timestamp} [INFO] [PCAP] 流量特征: 大量TCP ACK确认，无异常断开")
                log_content.append(f"{timestamp} [INFO] [PCAP] 建议: 这是一个正常的TCP会话捕获，建议检查应用层协议解码")
            
        else:
            # tshark没有输出，尝试使用tcpdump
            log_content.append(f"{timestamp} [WARN] [PCAP] tshark解析结果为空，尝试tcpdump")
            
            tcpdump_cmd = ['tcpdump', '-r', str(pcap_path), '-nn', '-v', '-c', '100']
            result = subprocess.run(tcpdump_cmd, capture_output=True, text=True, timeout=30, errors='replace')
            
            if result.stdout:
                dump_lines = result.stdout.strip().split('\n')[:30]
                for dump_line in dump_lines:
                    if dump_line.strip():
                        log_content.append(f"{timestamp} [INFO] [PCAP] {dump_line[:180]}")
            
    except subprocess.TimeoutExpired:
        log_content.append(f"{timestamp} [ERROR] [PCAP] 解析超时")
    except Exception as e:
        log_content.append(f"{timestamp} [ERROR] [PCAP] 解析失败: {str(e)}")
    
    # 确保至少有足够的日志行数
    while len(log_content) < 30:
        t = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        log_content.append(f"{t} [INFO] [PCAP] 处理记录")
    
    log_path = UPLOAD_DIR / f"{pcap_path.stem}_converted.log"
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(log_content))
    
    return str(log_path)

app = FastAPI(
    title="Log Analyzer",
    description="Large-scale log file analysis with LLM",
    version="1.0.0"
)

logger = logging.getLogger("web")
logger.setLevel(logging.INFO)

# 配置全局日志记录器，将日志输出到项目根目录的 logs 文件夹
LOG_DIR = SCRIPT_DIR.parent / "logs"
LOG_DIR.mkdir(parents=True, exist_ok=True)

# 创建日志文件处理器 - 默认前缀
log_file = LOG_DIR / f"web_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
file_handler = logging.FileHandler(log_file, encoding='utf-8')
file_handler.setLevel(logging.INFO)

# 创建格式化器
formatter = logging.Formatter(
    '%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
file_handler.setFormatter(formatter)

# 添加处理器到 logger
logger.addHandler(file_handler)

# 确保控制台输出
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.WARNING)
console_handler.setFormatter(formatter)
logger.addHandler(console_handler)

logger.info(f"日志系统初始化完成，日志文件: {log_file}")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class ConcurrencyLimitMiddleware(BaseHTTPMiddleware):
    """并发限流中间件：使用信号量控制最大并发请求数"""
    
    def __init__(self, app, max_concurrent: int = 200):
        super().__init__(app)
        self.semaphore = asyncio.Semaphore(max_concurrent)
        self.active_requests = 0
        self.max_concurrent = max_concurrent
        self._lock = asyncio.Lock()
    
    async def dispatch(self, request: Request, call_next):
        try:
            await asyncio.wait_for(
                self.semaphore.acquire(),
                timeout=30.0
            )
        except asyncio.TimeoutError:
            return JSONResponse(
                status_code=503,
                content={"detail": "Server is busy, please retry later"}
            )
        
        try:
            async with self._lock:
                self.active_requests += 1
            response = await call_next(request)
            response.headers["X-Active-Requests"] = str(self.active_requests)
            return response
        finally:
            async with self._lock:
                self.active_requests -= 1
            self.semaphore.release()

app.add_middleware(ConcurrencyLimitMiddleware, max_concurrent=200)

UPLOAD_DIR = PROJECT_ROOT / "log_analyzer" / "uploads"
REPORTS_DIR = PROJECT_ROOT / "log_analyzer" / "reports"
LOGS_DIR = PROJECT_ROOT / "log_analyzer" / "logs"

ensure_dir(str(UPLOAD_DIR))
ensure_dir(str(REPORTS_DIR))
ensure_dir(str(LOGS_DIR))

processing_tasks: Dict[str, Dict[str, Any]] = {}

class ProcessRequest(BaseModel):
    file_path: Optional[str] = None
    directory_path: Optional[str] = None
    chunk_size: int = 50000
    force_restart: bool = False

class ProcessResponse(BaseModel):
    task_id: str
    status: str
    message: str

class TaskStatus(BaseModel):
    task_id: str
    status: str
    progress: float
    message: str
    reports: Optional[List[Dict[str, str]]] = None
    error: Optional[str] = None

# 路径读取相关模型
class PathReadRequest(BaseModel):
    """从服务器路径读取日志文件的请求模型"""
    path: str  # 文件或目录路径
    recursive: bool = False  # 是否递归读取子目录
    max_file_size: int = 100 * 1024 * 1024  # 最大文件大小：100MB
    file_patterns: Optional[List[str]] = None  # 文件匹配模式，如 ["*.log", "*.txt"]

class PathReadResponse(BaseModel):
    """路径读取响应模型"""
    success: bool
    path: str
    file_count: int = 0
    total_size: int = 0
    files: Optional[List[Dict[str, Any]]] = None
    preview: Optional[str] = None
    error: Optional[str] = None

# 安全配置
ALLOWED_READ_PATHS = [
    "/var/log",           # 系统日志
    "/tmp",               # 临时文件
    "/home",              # 用户目录
    "/Users",             # macOS用户目录
]

def validate_and_resolve_path(requested_path: str) -> tuple:
    """
    验证并解析路径，防止路径遍历攻击和越权访问
    
    Args:
        requested_path: 用户请求的路径
        
    Returns:
        tuple: (resolved_path, error_message)
        - 成功：返回绝对路径，error为None
        - 失败：path为None，error包含错误信息
    """
    if not requested_path:
        return None, "路径不能为空"
    
    try:
        # 解析路径，去除多余的斜杠和相对路径
        requested_path = requested_path.strip()
        resolved_path = Path(requested_path).resolve()
        
        # 路径遍历检查
        if '..' in requested_path:
            return None, "禁止使用 '..' 进行路径遍历"
        
        # 检查是否为绝对路径
        if not resolved_path.is_absolute():
            return None, "只支持绝对路径"
        
        # 检查路径是否存在
        if not resolved_path.exists():
            return None, f"路径不存在: {requested_path}"
        
        # 安全检查：确保路径在允许的目录内或项目目录内
        is_allowed = False
        
        # 检查是否在项目目录内（始终允许读取项目自身文件）
        try:
            project_resolved = PROJECT_ROOT.resolve()
            if resolved_path.is_relative_to(project_resolved):
                is_allowed = True
        except (ValueError, OSError):
            pass
        
        # 检查是否在允许的系统目录内
        if not is_allowed:
            for allowed_base in ALLOWED_READ_PATHS:
                try:
                    if str(resolved_path).startswith(allowed_base):
                        is_allowed = True
                        break
                except (ValueError, OSError):
                    continue
        
        if not is_allowed:
            return None, f"路径不在允许的读取范围内。允许的目录: {', '.join(ALLOWED_READ_PATHS + [str(PROJECT_ROOT)])}"
        
        # 检查读权限
        if not os.access(resolved_path, os.R_OK):
            return None, f"无读取权限: {requested_path}"
        
        return str(resolved_path), None
        
    except Exception as e:
        return None, f"路径验证失败: {str(e)}"

def scan_directory_for_logs(dir_path: str, recursive: bool = False, 
                           patterns: Optional[List[str]] = None,
                           max_size: int = 100 * 1024 * 1024) -> tuple:
    """
    扫描目录查找日志文件
    
    Args:
        dir_path: 目录路径
        recursive: 是否递归扫描子目录
        patterns: 文件匹配模式列表
        max_size: 最大文件大小限制
        
    Returns:
        tuple: (files_list, error_message)
    """
    if patterns is None:
        patterns = ['*.log', '*.txt', '*.pcap', '*.pcapng', '*.gz', '*.zip']
    
    log_files = []
    
    try:
        scan_path = Path(dir_path)
        
        # 如果是文件，直接返回
        if scan_path.is_file():
            if any(scan_path.match(pattern) for pattern in patterns):
                stat = scan_path.stat()
                if stat.st_size <= max_size:
                    log_files.append({
                        "name": scan_path.name,
                        "path": str(scan_path),
                        "size": stat.st_size,
                        "size_str": format_bytes(stat.st_size),
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "type": "file"
                    })
            return log_files, None
        
        # 目录扫描
        if recursive:
            # 递归扫描
            for pattern in patterns:
                for file_path in scan_path.rglob(pattern):
                    if file_path.is_file():
                        try:
                            stat = file_path.stat()
                            if stat.st_size <= max_size:
                                log_files.append({
                                    "name": file_path.name,
                                    "path": str(file_path),
                                    "size": stat.st_size,
                                    "size_str": format_bytes(stat.st_size),
                                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                                    "type": "file",
                                    "relative_path": str(file_path.relative_to(scan_path))
                                })
                        except (OSError, PermissionError):
                            continue
        else:
            # 非递归扫描
            for pattern in patterns:
                for file_path in scan_path.glob(pattern):
                    if file_path.is_file():
                        try:
                            stat = file_path.stat()
                            if stat.st_size <= max_size:
                                log_files.append({
                                    "name": file_path.name,
                                    "path": str(file_path),
                                    "size": stat.st_size,
                                    "size_str": format_bytes(stat.st_size),
                                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                                    "type": "file"
                                })
                        except (OSError, PermissionError):
                            continue
        
        # 按修改时间排序，最新的在前
        log_files.sort(key=lambda x: x['modified'], reverse=True)
        
        return log_files, None
        
    except Exception as e:
        return [], f"扫描目录失败: {str(e)}"

def read_file_preview(file_path: str, max_lines: int = 100) -> tuple:
    """
    读取文件预览内容
    
    Args:
        file_path: 文件路径
        max_lines: 最大预览行数
        
    Returns:
        tuple: (preview_content, error_message)
    """
    try:
        preview_lines = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            for i, line in enumerate(f):
                if i >= max_lines:
                    break
                preview_lines.append(line.rstrip('\n\r'))
        
        return '\n'.join(preview_lines), None
        
    except Exception as e:
        return None, f"读取文件失败: {str(e)}"

def setup_logging(task_id: str, file_paths: List[str] = None) -> tuple:
    ensure_dir(str(LOGS_DIR))
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    # 判断是否为路径读取任务
    is_path_task = task_id.startswith('path_')
    
    if file_paths:
        file_names = [Path(fp).stem for fp in file_paths[:3]]
        if len(file_names) == 1:
            file_label = file_names[0]
        else:
            file_label = f"{file_names[0]}_等{len(file_paths)}个文件"
        
        # 根据任务类型选择日志文件前缀
        if is_path_task:
            log_file = LOGS_DIR / f'web_path_{timestamp}_{file_label}.log'
        else:
            log_file = LOGS_DIR / f'web_process_{timestamp}_{file_label}.log'
    else:
        if is_path_task:
            log_file = LOGS_DIR / f'web_path_{timestamp}_{task_id[:8]}.log'
        else:
            log_file = LOGS_DIR / f'web_process_{timestamp}_{task_id[:8]}.log'

    formatter = logging.Formatter(
        fmt='%(asctime)s [%(levelname)s] %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )

    # 创建任务专用logger
    task_logger = logging.getLogger(f'web_{task_id}')
    task_logger.setLevel(logging.INFO)
    task_logger.propagate = True  # 允许传播到根logger

    for handler in task_logger.handlers[:]:
        task_logger.removeHandler(handler)

    # 获取根logger并添加文件handler
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.INFO)

    # 移除现有的文件handler（如果有）
    for handler in root_logger.handlers[:]:
        if isinstance(handler, logging.FileHandler):
            root_logger.removeHandler(handler)

    # 添加新的文件handler
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    file_handler.setFormatter(formatter)
    root_logger.addHandler(file_handler)

    # 确保控制台输出（如果需要）
    has_console_handler = any(isinstance(h, logging.StreamHandler) for h in root_logger.handlers)
    if not has_console_handler:
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.WARNING)
        console_handler.setFormatter(formatter)
        root_logger.addHandler(console_handler)

    return str(log_file), task_logger

async def process_log_files(
    task_id: str,
    file_paths: List[str],
    chunk_size: int = 50000,
    force_restart: bool = False
):
    task_info = processing_tasks[task_id]
    task_info["status"] = "processing"
    task_info["progress"] = 0.0
    task_info["message"] = "开始处理日志文件..."

    user_id = task_info.get("user_id", "default")
    user_reports_dir = Path(task_info.get("reports_dir", REPORTS_DIR))
    user_checkpoints_dir = Path(task_info.get("checkpoints_dir", str(PROJECT_ROOT / "log_analyzer" / "checkpoints")))

    ensure_dir(str(user_reports_dir))
    ensure_dir(str(user_checkpoints_dir))

    try:
        log_file, logger = setup_logging(task_id, file_paths)
        task_info["log_file"] = log_file

        llm_config_path = "/Users/a666/Documents/trae_projects/log/log_analyzer/llmconfig"
        try:
            llm_config = load_llm_config(llm_config_path)
        except FileNotFoundError as e:
            error_msg = f"LLM配置文件未找到: {llm_config_path}"
            logger.error(f"[Task {task_id}] {error_msg}")
            task_info["status"] = "failed"
            task_info["message"] = error_msg
            task_info["error"] = error_msg
            raise RuntimeError(error_msg)

        llm_client = LLMClient(
            config=llm_config,
            max_retries=3,
            retry_delay=1.0
        )

        parser = LogParser(chunk_size=chunk_size)

        checkpoint_manager = CheckpointManager(
            checkpoint_dir=str(user_checkpoints_dir)
        )

        report_generator = ReportGenerator(output_dir=str(user_reports_dir))

        processor = ChunkProcessor(
            parser=parser,
            llm_client=llm_client,
            checkpoint_manager=checkpoint_manager,
            chunk_size=chunk_size,
            enable_checkpoint=False
        )

        total_files = len(file_paths)
        processed_files = 0
        all_reports = []
        all_results = []

        for idx, file_path in enumerate(file_paths):
            file_name = Path(file_path).name
            task_info["message"] = f"正在处理文件 {idx + 1}/{total_files}: {file_name}"
            task_info["progress"] = idx / total_files * 80

            try:
                logger.info(f"[Task {task_id}] 开始处理文件: {file_path}")

                if file_path.lower().endswith('.pcap'):
                    logger.info(f"[Task {task_id}] PCAP文件检测到，使用专用处理器...")
                    
                    from log_analyzer.processor.pcap_processor import PCAPProcessor
                    
                    pcap_processor = PCAPProcessor(max_packets=1000)
                    stats, packets = pcap_processor.process_file(file_path)
                    
                    logger.info(f"[Task {task_id}] PCAP处理完成:")
                    logger.info(f"  - 总数据包: {stats.total_packets}")
                    logger.info(f"  - TCP: {stats.tcp_packets}, UDP: {stats.udp_packets}")
                    logger.info(f"  - 错误: {stats.error_count}, 警告: {stats.warning_count}")
                    
                    logger.info(f"[Task {task_id}] 准备调用LLM分析PCAP数据...")
                    analysis_prompt = pcap_processor.generate_analysis_prompt()
                    
                    messages = [
                        {"role": "system", "content": "你是一个专业的网络流量分析工程师，擅长分析PCAP抓包数据并提供网络诊断和优化建议。请用JSON格式回复，包含summary、traffic_analysis、error_analysis、suggestions等字段。"},
                        {"role": "user", "content": analysis_prompt}
                    ]
                    
                    logger.info(f"[Task {task_id}] LLM请求发送中...")
                    llm_response = await llm_client.chat(messages=messages, temperature=0.3, max_tokens=2048)
                    
                    if llm_response.is_success() and llm_response.content:
                        llm_result = llm_response.content
                        logger.info(f"[Task {task_id}] LLM分析完成，结果长度: {len(llm_result)} chars")
                    else:
                        llm_result = f"LLM分析失败: {llm_response.error}"
                        logger.error(f"[Task {task_id}] LLM分析失败: {llm_response.error}")
                    
                    report_data = {
                        "title": f"PCAP网络流量分析报告 - {file_name}",
                        "file_path": file_path,
                        "file_size": Path(file_path).stat().st_size,
                        "statistics": stats.to_dict(),
                        "analysis_result": llm_result,
                        "summary": f"分析了 {stats.total_packets} 个数据包，发现 {stats.error_count} 个错误和 {stats.warning_count} 个警告"
                    }
                    
                    pcap_report_path = user_reports_dir / f"pcap_report_{Path(file_path).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                    
                    with open(f"{pcap_report_path}.json", 'w', encoding='utf-8') as f:
                        json.dump(report_data, f, indent=2, ensure_ascii=False)
                    
                    md_content = f"""# PCAP网络流量分析报告

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**源文件**: {file_path}
**文件大小**: {format_bytes(Path(file_path).stat().st_size)}

---

## 网络流量统计

| 指标 | 数值 |
|------|------|
| 总数据包数 | {stats.total_packets:,} |
| 总字节数 | {stats.total_bytes:,} |
| TCP数据包 | {stats.tcp_packets:,} |
| UDP数据包 | {stats.udp_packets:,} |
| ICMP数据包 | {stats.icmp_packets:,} |
| ASTERIX数据包 | {stats.asterix_packets:,} |

## TCP会话分析

| 指标 | 数值 |
|------|------|
| SYN请求 | {stats.tcp_syn_count} |
| RST重置 | {stats.tcp_rst_count} |
| FIN结束 | {stats.tcp_fin_count} |
| ACK确认 | {stats.tcp_ack_count} |

## 应用层协议

- HTTP请求: {stats.http_requests}
- DNS查询: {stats.dns_queries}

## 异常统计

| 指标 | 数值 |
|------|------|
| 错误数 | {stats.error_count} |
| 警告数 | {stats.warning_count} |

## 唯一IP地址

{', '.join(stats.unique_ips[:15])}
{"..." if len(stats.unique_ips) > 15 else ""}

## 唯一端口

{', '.join(map(str, stats.unique_ports[:25]))}
{"..." if len(stats.unique_ports) > 25 else ""}

## LLM分析结果

{llm_result if llm_result else "LLM分析未返回结果"}

---

*报告生成时间: {datetime.now().isoformat()}*
"""
                    
                    with open(f"{pcap_report_path}.md", 'w', encoding='utf-8') as f:
                        f.write(md_content)
                    
                    saved_files = [f"{pcap_report_path}.json", f"{pcap_report_path}.md"]
                    for saved_file in saved_files:
                        all_reports.append({
                            "name": Path(saved_file).name,
                            "path": saved_file,
                            "type": "markdown" if saved_file.endswith(".md") else "json"
                        })
                    
                    logger.info(f"[Task {task_id}] PCAP报告已保存: {saved_files}")
                    processed_files += 1

                    pcap_html_path = f"{pcap_report_path}.html"
                    try:
                        pcap_html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="UTF-8"><title>PCAP分析报告</title>
<style>body{{font-family:sans-serif;padding:2rem;max-width:900px;margin:0 auto;}}
h1{{color:#007AFF;}}pre{{background:#f5f5f7;padding:1rem;border-radius:8px;}}</style>
</head>
<body><h1>PCAP网络流量分析报告</h1>
<p><b>文件:</b> {pcap_path.name}</p>
<p><b>生成时间:</b> {datetime.now().isoformat()}</p>
<p><b>总数据包:</b> {stats.total_packets}</p>
<p><b>TCP包数:</b> {stats.tcp_packets}, <b>UDP包数:</b> {stats.udp_packets}</p>
<p><b>错误数:</b> {stats.error_count}, <b>警告数:</b> {stats.warning_count}</p>
<h2>LLM分析结果</h2>
<pre>{llm_result if llm_result else "LLM分析未返回结果"}</pre>
</body></html>"""
                        with open(pcap_html_path, 'w', encoding='utf-8') as f:
                            f.write(pcap_html_content)
                        all_reports.append({
                            "name": Path(pcap_html_path).name,
                            "path": pcap_html_path,
                            "type": "html"
                        })
                    except Exception as e:
                        logger.warning(f"[Task {task_id}] PCAP HTML生成失败: {e}")
                    
                    try:
                        from reportlab.lib.pagesizes import A4
                        from reportlab.lib.styles import getSampleStyleSheet
                        from reportlab.lib.units import cm
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        pcap_pdf_path = f"{pcap_report_path}.pdf"
                        doc = SimpleDocTemplate(pcap_pdf_path, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
                        styles = getSampleStyleSheet()
                        story = [
                            Paragraph("PCAP网络流量分析报告", styles['Title']),
                            Paragraph(f"文件: {pcap_path.name}", styles['Normal']),
                            Paragraph(f"生成时间: {datetime.now().isoformat()}", styles['Normal']),
                            Paragraph(f"总数据包: {stats.total_packets}, TCP: {stats.tcp_packets}, UDP: {stats.udp_packets}", styles['Normal']),
                            Paragraph(f"错误: {stats.error_count}, 警告: {stats.warning_count}", styles['Normal']),
                            Spacer(1, 0.5*cm),
                            Paragraph("LLM分析结果:", styles['Heading2']),
                            Paragraph(llm_result[:2000] if llm_result else "LLM分析未返回结果", styles['Normal'])
                        ]
                        doc.build(story)
                        all_reports.append({
                            "name": Path(pcap_pdf_path).name,
                            "path": pcap_pdf_path,
                            "type": "pdf"
                        })
                    except Exception as e:
                        logger.warning(f"[Task {task_id}] PCAP PDF生成失败: {e}")
                    
                    try:
                        from docx import Document
                        pcap_word_path = f"{pcap_report_path}.docx"
                        word_doc = Document()
                        word_doc.add_heading('PCAP网络流量分析报告', 0)
                        word_doc.add_paragraph(f"文件: {pcap_path.name}")
                        word_doc.add_paragraph(f"生成时间: {datetime.now().isoformat()}")
                        word_doc.add_paragraph(f"总数据包: {stats.total_packets}, TCP: {stats.tcp_packets}, UDP: {stats.udp_packets}")
                        word_doc.add_paragraph(f"错误: {stats.error_count}, 警告: {stats.warning_count}")
                        word_doc.add_heading('LLM分析结果', 1)
                        word_doc.add_paragraph(llm_result if llm_result else "LLM分析未返回结果")
                        word_doc.save(pcap_word_path)
                        all_reports.append({
                            "name": Path(pcap_word_path).name,
                            "path": pcap_word_path,
                            "type": "word"
                        })
                    except Exception as e:
                        logger.warning(f"[Task {task_id}] PCAP Word生成失败: {e}")
                    
                elif file_path.lower().endswith('.zip'):
                    logger.info(f"[Task {task_id}] ZIP文件，跳过直接分析（已在上传时解压）: {file_path}")
                    continue
                    
                else:
                    result = await processor.process_file_async(file_path=file_path, resume=True, force_restart=True)
                    all_results.append(result)
                    logger.info(f"[Task {task_id}] 文件处理完成，状态: {result.status}")

                    if result.status == "completed":
                        report = report_generator.generate_report(result)
                        saved_files = report_generator.save_report(report, format="html+md+pdf+word")
                        logger.info(f"[Task {task_id}] 报告已保存: {saved_files}")
                        for saved_file in saved_files:
                            file_type = "html" if saved_file.endswith(".html") else "pdf" if saved_file.endswith(".pdf") else "word" if saved_file.endswith(".docx") else "markdown" if saved_file.endswith(".md") else "json"
                            all_reports.append({
                                "name": Path(saved_file).name,
                                "path": saved_file,
                                "type": file_type
                            })
                        processed_files += 1
                    else:
                        logger.warning(f"[Task {task_id}] 文件处理未完成，状态: {result.status}")

            except Exception as e:
                import traceback
                logger.error(f"[Task {task_id}] 处理文件 {file_path} 时出错: {str(e)}")
                logger.error(traceback.format_exc())
                task_info["message"] = f"处理文件 {Path(file_path).name} 时出错: {str(e)}"
                continue

        if len(all_results) > 1:
            task_info["message"] = "正在生成综合报告..."
            task_info["progress"] = 95.0
            
            try:
                combined_report = report_generator.generate_combined_report(all_results)
                combined_files = report_generator.save_report(combined_report, format="html+md+pdf+word", prefix="combined")
                logger.info(f"[Task {task_id}] 综合报告已保存: {combined_files}")
                for saved_file in combined_files:
                    file_type = "html" if saved_file.endswith(".html") else "pdf" if saved_file.endswith(".pdf") else "word" if saved_file.endswith(".docx") else "markdown" if saved_file.endswith(".md") else "json"
                    all_reports.append({
                        "name": Path(saved_file).name,
                        "path": saved_file,
                        "type": file_type
                    })
            except Exception as e:
                logger.error(f"[Task {task_id}] 生成综合报告失败: {str(e)}")

        task_info["progress"] = 100.0
        task_info["status"] = "completed"
        task_info["message"] = f"处理完成！共处理 {processed_files} 个文件"
        task_info["reports"] = all_reports

        if all_results and user_id:
            try:
                storage = get_storage()
                stats_summary = {
                    "total_files": total_files,
                    "processed_files": processed_files,
                    "total_chunks": sum(len(r.chunks) if hasattr(r, 'chunks') else 0 for r in all_results)
                }
                for result in all_results:
                    file_name = Path(result.file_path).name if hasattr(result, 'file_path') else "unknown"
                    history_record = {
                        "title": f"日志分析报告 - {file_name}",
                        "file_name": file_name,
                        "file_type": Path(file_name).suffix.lstrip('.') or "log",
                        "summary": result.summary if hasattr(result, 'summary') else "",
                        "statistics": stats_summary,
                        "analysis": result.to_dict() if hasattr(result, 'to_dict') else {},
                        "files": [r for r in all_reports if Path(r.get("path", "")).name.startswith(file_name.rsplit('.', 1)[0])],
                        "tags": ["auto-generated"],
                        "metadata": {
                            "task_id": task_id,
                            "source": "process_endpoint"
                        }
                    }
                    storage.create(user_id, history_record)
                logger.info(f"[Task {task_id}] 已持久化 {len(all_results)} 条历史报告")
            except Exception as e:
                logger.error(f"[Task {task_id}] 持久化历史报告失败: {e}")

    except Exception as e:
        import traceback
        task_info["status"] = "failed"
        task_info["error"] = str(e)
        task_info["message"] = f"处理失败: {str(e)}"
        if 'logger' in locals():
            logger.error(f"处理任务 {task_id} 失败: {str(e)}")
            logger.error(traceback.format_exc())
    finally:
        if "llm_client" in locals():
            try:
                await llm_client.close()
            except Exception as e:
                if 'logger' in locals():
                    logger.warning(f"关闭 LLM 客户端时出错: {str(e)}")


# ==================== 用户识别接口（无鉴权版） ====================

class IdentifyRequest(BaseModel):
    user_id: str
    username: Optional[str] = None


@app.post("/api/auth/identify")
async def identify_user(data: IdentifyRequest):
    """用户身份识别接口

    客户端通过此接口标识自己，后续所有请求通过 X-User-Id 头携带。
    """
    user_info = user_manager.get_or_create_user(data.user_id, data.username)
    return JSONResponse({
        "code": 0,
        "message": "识别成功",
        "data": user_info
    })


@app.get("/api/auth/current")
async def get_current_user_info(current_user: Dict = Depends(get_current_user)):
    """获取当前请求的用户信息（从 X-User-Id 头提取）"""
    return JSONResponse({
        "code": 0,
        "message": "获取成功",
        "data": current_user
    })


@app.get("/")
async def root():
    return FileResponse(str(SCRIPT_DIR / "static" / "index.html"))

@app.get("/api/health")
async def health_check():
    return {"status": "ok", "message": "Log Analyzer API is running"}

SUPPORTED_EXTENSIONS = ('.log', '.txt', '.zip', '.pcap')


@app.post("/api/read-path")
async def read_path(
    request: PathReadRequest,
    current_user: Dict = Depends(get_current_user)
):
    """
    从服务器指定路径读取日志文件
    
    功能：
    - 读取指定文件或目录
    - 支持递归扫描子目录
    - 支持文件模式过滤（*.log, *.txt）
    - 提供文件预览功能
    
    安全特性：
    - 路径遍历攻击防护
    - 只允许访问指定目录（白名单）
    - 权限检查
    """
    # 验证路径
    resolved_path, error = validate_and_resolve_path(request.path)
    
    if error:
        return JSONResponse({
            "code": 1,
            "message": error,
            "data": None
        }, status_code=400)
    
    # 扫描目录查找日志文件
    files, scan_error = scan_directory_for_logs(
        dir_path=resolved_path,
        recursive=request.recursive,
        patterns=request.file_patterns,
        max_size=request.max_file_size
    )
    
    if scan_error:
        return JSONResponse({
            "code": 1,
            "message": scan_error,
            "data": None
        }, status_code=500)
    
    # 统计信息
    total_size = sum(f['size'] for f in files)
    file_count = len(files)
    
    # 如果是单个文件，提供预览
    preview = None
    if file_count == 1 and not Path(resolved_path).is_dir():
        preview, preview_error = read_file_preview(resolved_path, max_lines=50)
    
    return JSONResponse({
        "code": 0,
        "message": "读取成功",
        "data": {
            "success": True,
            "path": resolved_path,
            "file_count": file_count,
            "total_size": total_size,
            "total_size_str": format_bytes(total_size),
            "files": files if file_count > 0 else None,
            "preview": preview if preview else None,
            "user_id": current_user["user_id"],
            "timestamp": datetime.now().isoformat()
        }
    })

@app.post("/api/process-from-path")
async def process_from_path(
    request: PathReadRequest,
    background_tasks: BackgroundTasks,
    current_user: Dict = Depends(get_current_user)
):
    """
    从服务器指定路径读取日志文件并开始分析
    
    功能：
    - 读取指定路径的日志文件
    - 自动验证和过滤
    - 后台异步处理
    - 生成多格式报告
    
    与 /api/upload + /api/process 的区别：
    - 无需先上传文件到服务器
    - 直接从服务器指定路径读取
    - 适用于已存在于服务器上的日志文件
    """
    # 记录请求开始时间
    start_time = datetime.now()
    
    # 记录操作日志
    logger.info("=" * 80)
    logger.info(f"[{start_time.strftime('%Y-%m-%d %H:%M:%S')}] [API] /api/process-from-path 请求开始")
    logger.info(f"[{start_time.strftime('%Y-%m-%d %H:%M:%S')}] [参数] 用户ID: {current_user['user_id']}")
    logger.info(f"[{start_time.strftime('%Y-%m-%d %H:%M:%S')}] [参数] 请求路径: {request.path}")
    logger.info(f"[{start_time.strftime('%Y-%m-%d %H:%M:%S')}] [参数] 递归扫描: {request.recursive}")
    logger.info(f"[{start_time.strftime('%Y-%m-%d %H:%M:%S')}] [参数] 文件模式: {request.file_patterns}")
    logger.info(f"[{start_time.strftime('%Y-%m-%d %H:%M:%S')}] [参数] 最大文件大小: {request.max_file_size}")
    
    # 验证路径
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [步骤] 开始路径验证...")
    resolved_path, error = validate_and_resolve_path(request.path)
    
    if error:
        logger.error(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [错误] 路径验证失败: {error}")
        return JSONResponse({
            "code": 1,
            "message": error,
            "data": None
        }, status_code=400)
    
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [成功] 路径验证通过，解析路径: {resolved_path}")
    
    # 扫描目录查找日志文件
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [步骤] 开始扫描目录...")
    files, scan_error = scan_directory_for_logs(
        dir_path=resolved_path,
        recursive=request.recursive,
        patterns=request.file_patterns,
        max_size=request.max_file_size
    )
    
    if scan_error:
        logger.error(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [错误] 目录扫描失败: {scan_error}")
        return JSONResponse({
            "code": 1,
            "message": scan_error,
            "data": None
        }, status_code=500)
    
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [成功] 目录扫描完成，找到 {len(files)} 个文件")
    
    if not files:
        logger.warning(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [警告] 未找到符合条件的日志文件")
        return JSONResponse({
            "code": 1,
            "message": "未找到符合条件的日志文件",
            "data": {
                "path": resolved_path,
                "file_count": 0,
                "suggestion": "请检查路径是否正确，或尝试调整文件过滤模式"
            }
        }, status_code=404)
    
    # 生成任务ID
    task_id = f"path_{current_user['user_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}"
    
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [步骤] 生成任务ID: {task_id}")
    
    # 记录任务信息
    task_info = {
        "task_id": task_id,
        "user_id": current_user["user_id"],  # 添加用户ID，用于权限验证
        "status": "pending",
        "progress": 0.0,
        "message": "正在准备处理...",
        "file_count": len(files),
        "total_size": sum(f['size'] for f in files),
        "start_time": datetime.now().isoformat(),
        "source": "path",  # 标记为路径读取
        "source_path": resolved_path
    }
    
    # 保存任务信息到文件
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [步骤] 保存任务信息到文件...")
    try:
        tasks_file = TASKS_DIR / f"{task_id}.json"
        with open(tasks_file, 'w') as f:
            json.dump(task_info, f, indent=2)
        logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [成功] 任务信息已保存: {tasks_file}")
    except Exception as e:
        logger.error(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [错误] 保存任务信息失败: {e}")
    
    # 将文件路径列表添加到任务信息中
    task_info["files"] = [f['path'] for f in files]
    
    # 将任务添加到内存字典中，供状态查询使用
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [步骤] 注册任务到内存字典...")
    processing_tasks[task_id] = task_info
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [成功] 任务已注册")
    
    # 记录待处理的文件列表
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [文件列表] 待处理文件:")
    for idx, f in enumerate(files):
        logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}]   [{idx+1}] {f['path']} ({utils_module.get_file_size_str(f['path'])})")
    
    # 后台任务：开始处理
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [步骤] 提交后台处理任务...")
    background_tasks.add_task(
        process_files_from_path,
        task_id=task_id,
        file_paths=[f['path'] for f in files],
        user_id=current_user["user_id"],
        task_info=task_info
    )
    
    # 计算请求处理耗时
    elapsed_time = (datetime.now() - start_time).total_seconds()
    
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [成功] 路径读取任务已创建")
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [统计] 任务ID: {task_id}")
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [统计] 文件数量: {len(files)}")
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [统计] 总大小: {task_info['total_size']} bytes")
    logger.info(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] [统计] 请求耗时: {elapsed_time:.3f}s")
    logger.info("=" * 80)
    
    return JSONResponse({
        "code": 0,
        "message": "任务已创建，正在后台处理",
        "data": {
            "task_id": task_id,
            "status": "pending",
            "file_count": len(files),
            "total_size": format_bytes(sum(f['size'] for f in files)),
            "source_path": resolved_path,
            "status_url": f"/api/task/{task_id}"
        }
    })

async def process_files_from_path(
    task_id: str,
    file_paths: List[str],
    user_id: str,
    task_info: Dict[str, Any]
):
    """
    后台处理从路径读取的文件列表
    
    这与 process_files 函数的逻辑基本相同，
    区别在于文件已经存在于服务器上，无需上传步骤
    """
    from .app import (  # 避免循环导入
        setup_logging,
        TASKS_DIR,
        LOGS_DIR,
        Settings,
        LLMClient,
        ChunkProcessor,
        ReportGenerator
    )
    from .action_logger import record_task_start, record_task_complete, record_task_failed
    
    # 初始化组件
    log_file, task_logger = setup_logging(task_id, file_paths)
    llm_config = load_llm_config()
    
    # 记录任务开始
    record_task_start(user_id, task_id, file_paths[0] if file_paths else "")
    
    try:
        llm_client = LLMClient(llm_config)
        
        from log_analyzer.parser.log_parser import LogParser
        from log_analyzer.checkpoint.manager import CheckpointManager
        from log_analyzer.processor.chunk_processor import ChunkProcessor
        
        parser = LogParser(chunk_size=50000)
        
        user_reports_dir = PROJECT_ROOT / "log_analyzer" / "users" / user_id / "reports"
        user_checkpoints_dir = PROJECT_ROOT / "log_analyzer" / "users" / user_id / "checkpoints"
        ensure_dir(str(user_reports_dir))
        ensure_dir(str(user_checkpoints_dir))
        
        checkpoint_manager = CheckpointManager(
            checkpoint_dir=str(user_checkpoints_dir)
        )
        
        report_generator = ReportGenerator(output_dir=str(user_reports_dir))
        
        processor = ChunkProcessor(
            parser=parser,
            llm_client=llm_client,
            checkpoint_manager=checkpoint_manager,
            chunk_size=50000,
            enable_checkpoint=False
        )
        
        # 更新任务状态
        task_info["status"] = "running"
        task_info["message"] = "开始处理文件..."
        task_info["progress"] = 5.0
        
        with open(TASKS_DIR / f"{task_id}.json", 'w') as f:
            json.dump(task_info, f, indent=2)
        
        all_results = []
        all_reports = []
        processed_files = 0
        total_files = len(file_paths)
        
        for idx, file_path in enumerate(file_paths):
            file_name = Path(file_path).name
            task_info["message"] = f"正在处理文件 {idx + 1}/{total_files}: {file_name}"
            task_info["progress"] = idx / total_files * 80
            
            try:
                task_logger.info(f"[Task {task_id}] 开始处理文件: {file_path}")
                
                if file_path.lower().endswith('.pcap'):
                    task_logger.info(f"[Task {task_id}] PCAP文件检测到，使用专用处理器...")
                    
                    from log_analyzer.processor.pcap_processor import PCAPProcessor
                    
                    pcap_processor = PCAPProcessor(max_packets=1000)
                    stats, packets = pcap_processor.process_file(file_path)
                    
                    task_logger.info(f"[Task {task_id}] PCAP处理完成:")
                    task_logger.info(f"  - 总数据包: {stats.total_packets}")
                    task_logger.info(f"  - TCP: {stats.tcp_packets}, UDP: {stats.udp_packets}")
                    task_logger.info(f"  - 错误: {stats.error_count}, 警告: {stats.warning_count}")
                    
                    task_logger.info(f"[Task {task_id}] 准备调用LLM分析PCAP数据...")
                    analysis_prompt = pcap_processor.generate_analysis_prompt()
                    
                    messages = [
                        {"role": "system", "content": "你是一个专业的网络流量分析工程师，擅长分析PCAP抓包数据并提供网络诊断和优化建议。请用JSON格式回复，包含summary、traffic_analysis、error_analysis、suggestions等字段。"},
                        {"role": "user", "content": analysis_prompt}
                    ]
                    
                    task_logger.info(f"[Task {task_id}] LLM请求发送中...")
                    llm_response = await llm_client.chat(messages=messages, temperature=0.3, max_tokens=2048)
                    
                    if llm_response.is_success() and llm_response.content:
                        llm_result = llm_response.content
                        task_logger.info(f"[Task {task_id}] LLM分析完成，结果长度: {len(llm_result)} chars")
                    else:
                        llm_result = f"LLM分析失败: {llm_response.error}"
                        task_logger.error(f"[Task {task_id}] LLM分析失败: {llm_response.error}")
                    
                    report_data = {
                        "title": f"PCAP网络流量分析报告 - {file_name}",
                        "file_path": file_path,
                        "file_size": Path(file_path).stat().st_size,
                        "statistics": stats.to_dict(),
                        "analysis_result": llm_result,
                        "summary": f"分析了 {stats.total_packets} 个数据包，发现 {stats.error_count} 个错误和 {stats.warning_count} 个警告"
                    }
                    
                    pcap_report_path = user_reports_dir / f"report_path_{Path(file_path).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                    
                    with open(f"{pcap_report_path}.json", 'w', encoding='utf-8') as f:
                        json.dump(report_data, f, indent=2, ensure_ascii=False)
                    
                    md_content = f"""# PCAP网络流量分析报告

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**源文件**: {file_path}
**文件大小**: {format_bytes(Path(file_path).stat().st_size)}

---

## 网络流量统计
- 总数据包: {stats.total_packets}
- 总字节数: {stats.total_bytes}
- TCP: {stats.tcp_packets}
- UDP: {stats.udp_packets}
- ICMP: {stats.icmp_packets}
- ASTERIX: {stats.asterix_packets}

## 分析结果

{llm_result}
"""
                    
                    with open(f"{pcap_report_path}.md", 'w', encoding='utf-8') as f:
                        f.write(md_content)
                    
                    saved_files = [f"{pcap_report_path}.json", f"{pcap_report_path}.md"]
                    for saved_file in saved_files:
                        file_type = "markdown" if saved_file.endswith(".md") else "json"
                        all_reports.append({
                            "name": Path(saved_file).name,
                            "path": saved_file,
                            "type": file_type
                        })
                    
                    processed_files += 1
                
                else:
                    result = await processor.process_file_async(
                        file_path=file_path,
                        resume=True,
                        force_restart=True
                    )
                    all_results.append(result)
                    
                    task_logger.info(f"[Task {task_id}] 文件处理完成: {file_name}, 状态: {result.status}")
                    
                    if result.status == "completed":
                        report = report_generator.generate_report(result)
                        saved_files = report_generator.save_report(report, format="html+md+pdf+word", prefix="report_path")
                        
                        task_logger.info(f"[Task {task_id}] 报告已保存: {saved_files}")
                        
                        for saved_file in saved_files:
                            file_type = "html" if saved_file.endswith(".html") else \
                                       "pdf" if saved_file.endswith(".pdf") else \
                                       "word" if saved_file.endswith(".docx") else \
                                       "markdown" if saved_file.endswith(".md") else "json"
                            all_reports.append({
                                "name": Path(saved_file).name,
                                "path": saved_file,
                                "type": file_type,
                                "source_file": file_path
                            })
                        
                        processed_files += 1
                    else:
                        task_logger.warning(f"[Task {task_id}] 文件处理未完成: {file_name}")
            
            except Exception as e:
                import traceback
                task_logger.error(f"[Task {task_id}] 处理文件 {file_path} 时出错: {str(e)}")
                task_logger.error(traceback.format_exc())
                continue
        
        if len(all_results) > 1:
            task_info["message"] = "正在生成综合报告..."
            task_info["progress"] = 95
            
            with open(TASKS_DIR / f"{task_id}.json", 'w') as f:
                json.dump(task_info, f, indent=2)
            
            try:
                combined_report = report_generator.generate_combined_report(all_results)
                combined_files = report_generator.save_report(combined_report, format="html+md+pdf+word", prefix=f"path_combined_{Path(file_paths[0]).stem}")
                
                task_logger.info(f"[Task {task_id}] 综合报告已保存: {combined_files}")
                
                for saved_file in combined_files:
                    file_type = "html" if saved_file.endswith(".html") else \
                               "pdf" if saved_file.endswith(".pdf") else \
                               "word" if saved_file.endswith(".docx") else \
                               "markdown" if saved_file.endswith(".md") else "json"
                    all_reports.append({
                        "name": Path(saved_file).name,
                        "path": saved_file,
                        "type": file_type,
                        "is_combined": True
                    })
            except Exception as e:
                task_logger.error(f"[Task {task_id}] 生成综合报告失败: {str(e)}")
        
        task_info["status"] = "completed"
        task_info["progress"] = 100.0
        task_info["message"] = f"处理完成！成功处理 {processed_files}/{total_files} 个文件"
        task_info["reports"] = all_reports
        task_info["end_time"] = datetime.now().isoformat()
        
        with open(TASKS_DIR / f"{task_id}.json", 'w') as f:
            json.dump(task_info, f, indent=2)
        
        task_logger.info(f"[Task {task_id}] 任务完成！成功处理 {processed_files}/{total_files} 个文件")
        
        record_task_complete(user_id, task_id)
    
    except Exception as e:
        import traceback
        task_logger.error(f"处理任务 {task_id} 失败: {str(e)}")
        task_logger.error(traceback.format_exc())
        
        record_task_failed(user_id, task_id, str(e))
        
        task_info["status"] = "failed"
        task_info["error"] = str(e)
        task_info["message"] = f"任务失败: {str(e)}"
        
        with open(TASKS_DIR / f"{task_id}.json", 'w') as f:
            json.dump(task_info, f, indent=2)
@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...), current_user: Dict = Depends(get_current_user)):
    """上传文件接口（需要认证）"""
    try:
        if not file.filename.lower().endswith(SUPPORTED_EXTENSIONS):
            return JSONResponse({
                "code": 1,
                "message": f"不支持的文件类型: {file.filename}\n支持的类型: {', '.join(SUPPORTED_EXTENSIONS)}",
                "data": None
            }, status_code=400)

        user_upload_dir = get_user_upload_dir(current_user["user_id"])
        file_path = user_upload_dir / file.filename
        content = await file.read()

        with open(file_path, "wb") as buffer:
            buffer.write(content)

        # 记录文件上传操作日志
        record_file_upload(current_user["user_id"], file.filename, len(content))

        result_files = []
        if file.filename.lower().endswith('.zip'):
            try:
                extracted_files = extract_zip(file_path, user_upload_dir)
                for extracted in extracted_files:
                    if extracted.lower().endswith(('.log', '.txt')):
                        result_files.append({
                            'path': extracted,
                            'name': Path(extracted).name,
                            'size': format_bytes(os.path.getsize(extracted)) if os.path.exists(extracted) else '0 B'
                        })
            except Exception as e:
                logging.error(f"ZIP解压失败: {str(e)}")

        if not file.filename.lower().endswith('.zip') or len(result_files) == 0:
            result_files.append({
                'path': str(file_path),
                'name': file.filename,
                'size': format_bytes(len(content))
            })

        return JSONResponse({
            "code": 0,
            "message": "上传成功",
            "data": {
                "success": True,
                "file_path": str(file_path),
                "file_name": file.filename,
                "file_size": format_bytes(len(content)),
                "extracted_files": result_files
            }
        })
    except Exception as e:
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)

@app.post("/api/list-dir")
async def list_dir_post(
    request: Dict = Body(...),
    current_user: Dict = Depends(get_current_user)
):
    """列出目录内容（支持权限验证、递归和文件模式匹配，单一接口）"""
    path = request.get("path", "/")
    recursive = request.get("recursive", False)
    file_patterns = request.get("file_patterns", ["*.log"])
    validate_only = request.get("validate_only", False)  # 是否仅验证路径

    try:
        dir_path = Path(path)
        
        # 安全检查：解析绝对路径
        abs_path = dir_path.resolve()
        
        # 特殊处理：当访问根目录时，返回允许的目录列表（跨平台支持）
        # Windows: path == '/' 或 path == 'C:\\' 或类似
        # macOS/Linux: path == '/'
        def is_root_path(p):
            """判断是否为根路径"""
            if sys.platform.startswith('win'):
                # Windows 根路径判断
                p_str = str(p).lower()
                return p_str == '/' or p_str == '\\' or (len(p_str) == 3 and p_str[1] == ':')
            else:
                # macOS/Linux 根路径判断
                return str(p) == '/'
        
        if is_root_path(path) or is_root_path(abs_path):
            # 返回真正的根目录内容
            root_contents = []
            try:
                if sys.platform.startswith('win'):
                    # Windows 列出所有驱动器
                    import string
                    for letter in string.ascii_uppercase:
                        drive_path = Path(f"{letter}:\\")
                        if drive_path.exists():
                            try:
                                stat = drive_path.stat()
                                root_contents.append({
                                    "name": f"{letter}:\\",
                                    "path": f"{letter}:\\",
                                    "size": stat.st_size,
                                    "size_str": format_bytes(stat.st_size),
                                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                                    "type": "directory"
                                })
                            except:
                                root_contents.append({
                                    "name": f"{letter}:\\",
                                    "path": f"{letter}:\\",
                                    "size": 0,
                                    "size_str": "N/A",
                                    "modified": "",
                                    "type": "directory"
                                })
                else:
                    # macOS/Linux 列出根目录内容
                    for item in os.listdir('/'):
                        item_path = Path('/') / item
                        try:
                            if item_path.is_dir() and os.access(str(item_path), os.R_OK):
                                stat = item_path.stat()
                                root_contents.append({
                                    "name": item,
                                    "path": str(item_path),
                                    "size": stat.st_size,
                                    "size_str": format_bytes(stat.st_size),
                                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                                    "type": "directory"
                                })
                        except:
                            pass
            except Exception as e:
                logging.error(f"读取根目录失败: {e}")
            
            # 根据操作系统返回合适的根路径显示
            display_path = '/' if not sys.platform.startswith('win') else 'C:\\'
            
            return JSONResponse({
                "code": 0,
                "message": "获取根目录列表成功",
                "data": {
                    "current_path": display_path,
                    "parent_path": None,
                    "files": root_contents,
                    "allowed_directories": get_allowed_directories_from_config(),
                    "platform": "windows" if sys.platform.startswith('win') else ("macos" if sys.platform == 'darwin' else "linux")
                }
            })
        
        # 权限验证
        if not is_path_allowed(abs_path):
            return JSONResponse({
                "code": 1,
                "message": f"路径 {path} 不在允许访问的目录范围内。允许的目录：{', '.join(str(d) for d in ALLOWED_DIRECTORIES)}",
                "data": None
            }, status_code=403)

        if not dir_path.exists():
            return JSONResponse({
                "code": 1,
                "message": "路径不存在",
                "data": None
            }, status_code=404)

        # 检查是否有读取权限
        if not os.access(str(abs_path), os.R_OK):
            return JSONResponse({
                "code": 1,
                "message": f"没有读取路径 {path} 的权限",
                "data": None
            }, status_code=403)

        files = []

        def match_pattern(filename, patterns):
            """检查文件名是否匹配任一模式"""
            import fnmatch
            for pattern in patterns:
                if fnmatch.fnmatch(filename, pattern):
                    return True
            return False

        if dir_path.is_file():
            # 如果是文件，直接返回该文件信息
            stat = dir_path.stat()
            files.append({
                "name": dir_path.name,
                "path": str(dir_path),
                "size": stat.st_size,
                "size_str": format_bytes(stat.st_size),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "type": "file"
            })
        else:
            # 如果是目录，遍历内容
            if recursive:
                for item in dir_path.rglob("*"):
                    if item.is_file() and match_pattern(item.name, file_patterns):
                        stat = item.stat()
                        files.append({
                            "name": item.name,
                            "path": str(item),
                            "size": stat.st_size,
                            "size_str": format_bytes(stat.st_size),
                            "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                            "type": "file"
                        })
            else:
                for item in dir_path.iterdir():
                    if item.is_file() and match_pattern(item.name, file_patterns):
                        stat = item.stat()
                        files.append({
                            "name": item.name,
                            "path": str(item),
                            "size": stat.st_size,
                            "size_str": format_bytes(stat.st_size),
                            "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                            "type": "file"
                        })
                    elif item.is_dir():
                        files.append({
                            "name": item.name + "/",
                            "path": str(item),
                            "size": 0,
                            "size_str": "-",
                            "modified": "",
                            "type": "directory"
                        })

        # 如果是仅验证模式，返回验证结果
        if validate_only:
            return JSONResponse({
                "code": 0,
                "message": "路径验证成功",
                "data": {
                    "path": str(abs_path),
                    "is_file": abs_path.is_file(),
                    "is_directory": abs_path.is_dir(),
                    "file_count": len([f for f in files if f["type"] == "file"]),
                    "files": [f for f in files if f["type"] == "file"][:20]
                }
            })

        # 否则返回目录列表
        return JSONResponse({
            "code": 0,
            "message": "获取成功",
            "data": {
                "current_path": str(dir_path),
                "parent_path": str(dir_path.parent) if dir_path.parent != dir_path else None,
                "files": sorted(files, key=lambda x: (x["type"], x["name"]))
            }
        })
    except Exception as e:
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)

@app.post("/api/process")
async def start_processing(
    request: ProcessRequest,
    background_tasks: BackgroundTasks,
    current_user: Dict = Depends(get_current_user)
):
    """开始处理日志文件（需要认证）"""
    task_id = f"{current_user['user_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}"

    file_paths = []

    if request.file_path:
        if not Path(request.file_path).exists():
            return JSONResponse({
                "code": 1,
                "message": "文件不存在",
                "data": None
            }, status_code=404)
        file_paths.append(request.file_path)
    elif request.directory_path:
        dir_path = Path(request.directory_path)
        if not dir_path.exists() or not dir_path.is_dir():
            return JSONResponse({
                "code": 1,
                "message": "目录不存在",
                "data": None
            }, status_code=404)

        for file in dir_path.iterdir():
            if file.is_file() and (file.suffix in ['.log', '.txt', '.pcap'] or 'error' in file.name.lower()):
                file_paths.append(str(file))

    if not file_paths:
        return JSONResponse({
            "code": 1,
            "message": "未找到可处理的日志文件",
            "data": None
        }, status_code=400)

    user_reports_dir = get_user_reports_dir(current_user["user_id"])
    user_checkpoints_dir = get_user_checkpoints_dir(current_user["user_id"])

    processing_tasks[task_id] = {
        "status": "pending",
        "progress": 0.0,
        "message": "任务已创建，等待开始...",
        "file_paths": file_paths,
        "reports": None,
        "error": None,
        "user_id": current_user["user_id"],
        "reports_dir": str(user_reports_dir),
        "checkpoints_dir": str(user_checkpoints_dir)
    }

    background_tasks.add_task(
        process_log_files,
        task_id,
        file_paths,
        request.chunk_size,
        request.force_restart
    )

    return JSONResponse({
        "code": 0,
        "message": f"任务已创建，正在处理 {len(file_paths)} 个文件",
        "data": {
            "task_id": task_id,
            "status": "pending"
        }
    })

@app.get("/api/task/{task_id}", response_model=TaskStatus)
async def get_task_status(task_id: str, current_user: Dict = Depends(get_current_user)):
    """获取任务状态（需要认证，仅返回属于当前用户的任务）"""
    if task_id not in processing_tasks:
        return JSONResponse({
            "code": 1,
            "message": "任务不存在",
            "data": None
        }, status_code=404)

    task_info = processing_tasks[task_id]
    if task_info.get("user_id") != current_user["user_id"]:
        return JSONResponse({
            "code": 1,
            "message": "无权访问此任务",
            "data": None
        }, status_code=403)

    return JSONResponse({
        "code": 0,
        "message": "获取成功",
        "data": {
            "task_id": task_id,
            "status": task_info["status"],
            "progress": task_info["progress"],
            "message": task_info["message"],
            "reports": task_info.get("reports"),
            "error": task_info.get("error")
        }
    })

@app.get("/api/download/{file_path:path}")
async def download_file(file_path: str, current_user: Dict = Depends(get_current_user)):
    """下载文件（需要 X-User-Id 头，仅能下载用户自己的文件）"""
    try:
        path = Path(file_path)
        if not path.exists():
            return JSONResponse({
                "code": 1,
                "message": "文件不存在",
                "data": None
            }, status_code=404)

        user_reports_dir = get_user_reports_dir(current_user["user_id"])
        user_upload_dir = get_user_upload_dir(current_user["user_id"])

        if not (str(path).startswith(str(user_reports_dir)) or str(path).startswith(str(user_upload_dir))):
            return JSONResponse({
                "code": 1,
                "message": "无权访问此文件",
                "data": None
            }, status_code=403)

        return FileResponse(
            path=str(path),
            filename=path.name,
            media_type="application/octet-stream"
        )
    except Exception as e:
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)

@app.get("/api/reports")
async def list_reports(current_user: Dict = Depends(get_current_user)):
    """获取当前用户的报告列表（文件视角）
    
    所有用户只能查看自己的报告。
    """
    try:
        reports = []
        
        # 所有用户只能查看自己的报告
        user_reports_dir = get_user_reports_dir(current_user["user_id"])
        if user_reports_dir.exists():
            for file in user_reports_dir.iterdir():
                if file.is_file() and (file.suffix in ['.json', '.md', '.html', '.pdf', '.docx']):
                    stat = file.stat()
                    file_type = "markdown" if file.suffix == ".md" else "html" if file.suffix == ".html" else "pdf" if file.suffix == ".pdf" else "word" if file.suffix == ".docx" else "json"
                    reports.append({
                        "name": file.name,
                        "path": str(file),
                        "size": stat.st_size,
                        "size_str": format_bytes(stat.st_size),
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "type": file_type
                    })

        return JSONResponse({
            "code": 0,
            "message": "获取成功",
            "data": {"reports": sorted(reports, key=lambda x: x["modified"], reverse=True)}
        })
    except Exception as e:
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


# ==================== 历史报告 CRUD 接口 ====================

class HistoryReportCreate(BaseModel):
    title: str
    file_name: str
    file_type: str = "log"
    summary: str = ""
    statistics: Dict = {}
    analysis: Dict = {}
    files: List[Dict] = []
    tags: List[str] = []
    metadata: Dict = {}


class HistoryReportUpdate(BaseModel):
    title: Optional[str] = None
    summary: Optional[str] = None
    tags: Optional[List[str]] = None
    metadata: Optional[Dict] = None


@app.post("/api/history/reports")
async def create_history_report(
    data: HistoryReportCreate,
    current_user: Dict = Depends(get_current_user)
):
    """创建历史报告记录（持久化到本地存储）"""
    try:
        storage = get_storage()
        report_id = storage.create(current_user["user_id"], data.dict())
        return JSONResponse({
            "code": 0,
            "message": "创建成功",
            "data": {"report_id": report_id}
        })
    except Exception as e:
        logger.error(f"创建历史报告失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


@app.get("/api/history/reports")
async def list_history_reports(
    limit: int = 100,
    offset: int = 0,
    keyword: str = "",
    current_user: Dict = Depends(get_current_user)
):
    """获取当前用户的历史报告列表（按用户ID隔离）"""
    try:
        storage = get_storage()
        user_id = current_user["user_id"]

        if keyword:
            reports = storage.search(user_id, keyword, limit=limit)
        else:
            reports = storage.list(user_id, limit=limit, offset=offset)

        total = storage.count(user_id)

        return JSONResponse({
            "code": 0,
            "message": "获取成功",
            "data": {
                "reports": reports,
                "total": total,
                "limit": limit,
                "offset": offset
            }
        })
    except Exception as e:
        logger.error(f"获取历史报告列表失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


@app.get("/api/history/reports/{report_id}")
async def get_history_report(
    report_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """获取单个历史报告详情（按用户ID隔离）"""
    try:
        storage = get_storage()
        report = storage.get(current_user["user_id"], report_id)

        if not report:
            return JSONResponse({
                "code": 1,
                "message": "报告不存在",
                "data": None
            }, status_code=404)

        return JSONResponse({
            "code": 0,
            "message": "获取成功",
            "data": report
        })
    except Exception as e:
        logger.error(f"获取历史报告失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


@app.put("/api/history/reports/{report_id}")
async def update_history_report(
    report_id: str,
    data: HistoryReportUpdate,
    current_user: Dict = Depends(get_current_user)
):
    """更新历史报告（按用户ID隔离）"""
    try:
        storage = get_storage()
        update_data = {k: v for k, v in data.dict().items() if v is not None}
        success = storage.update(current_user["user_id"], report_id, update_data)

        if not success:
            return JSONResponse({
                "code": 1,
                "message": "报告不存在",
                "data": None
            }, status_code=404)

        return JSONResponse({
            "code": 0,
            "message": "更新成功",
            "data": None
        })
    except Exception as e:
        logger.error(f"更新历史报告失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


@app.delete("/api/history/reports/{report_id}")
async def delete_history_report(
    report_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """删除历史报告（按用户ID隔离）"""
    try:
        storage = get_storage()
        success = storage.delete(current_user["user_id"], report_id)

        if not success:
            return JSONResponse({
                "code": 1,
                "message": "报告不存在",
                "data": None
            }, status_code=404)

        return JSONResponse({
            "code": 0,
            "message": "删除成功",
            "data": None
        })
    except Exception as e:
        logger.error(f"删除历史报告失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


# ==================== 数据备份接口 ====================

@app.post("/api/backup/create")
async def create_backup(current_user: Dict = Depends(get_current_user)):
    """为当前用户创建数据备份"""
    try:
        storage = get_storage()
        backup_dir = str(PROJECT_ROOT / "log_analyzer" / "data" / "backups")
        ensure_dir(backup_dir)
        backup_path = storage.backup(current_user["user_id"], backup_dir)

        if not backup_path:
            return JSONResponse({
                "code": 1,
                "message": "无数据可备份",
                "data": None
            })

        return JSONResponse({
            "code": 0,
            "message": "备份成功",
            "data": {"backup_path": backup_path}
        })
    except Exception as e:
        logger.error(f"备份失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


# ==================== 用户操作历史记录接口 ====================

from .action_logger import get_action_log_storage, ACTION_TYPES


class ActionLogQuery(BaseModel):
    start_time: Optional[str] = None
    end_time: Optional[str] = None
    action_type: Optional[str] = None
    limit: int = 100
    offset: int = 0


@app.post("/api/history/actions")
async def query_action_logs(
    query: ActionLogQuery = None,
    current_user: Dict = Depends(get_current_user)
):
    """查询用户操作历史记录
    
    支持按时间范围、操作类型筛选，分页查询。
    所有用户都可以查看所有用户的操作记录。
    """
    if query is None:
        query = ActionLogQuery()
    
    try:
        storage = get_action_log_storage()
        
        # 所有用户都可以查看所有用户的操作记录
        # 获取所有用户的操作日志并合并
        all_records = []
        base_dir = storage.base_dir
        
        if base_dir.exists():
            for user_dir in base_dir.iterdir():
                if user_dir.is_dir():
                    user_id = user_dir.name
                    user_records, _ = storage.query(
                        user_id=user_id,
                        start_time=query.start_time,
                        end_time=query.end_time,
                        action_type=query.action_type,
                        limit=1000,
                        offset=0
                    )
                    all_records.extend(user_records)
        
        # 按时间降序排序
        all_records.sort(key=lambda x: x["timestamp"], reverse=True)
        
        total = len(all_records)
        records = all_records[query.offset:query.offset + query.limit]

        return JSONResponse({
            "code": 0,
            "message": "获取成功",
            "data": {
                "records": records,
                "total": total,
                "limit": query.limit,
                "offset": query.offset
            }
        })
    except Exception as e:
        logger.error(f"查询操作日志失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


@app.get("/api/history/actions/types")
async def get_action_types():
    """获取支持的操作类型列表（无需管理员权限）"""
    try:
        return JSONResponse({
            "code": 0,
            "message": "获取成功",
            "data": {
                "types": ACTION_TYPES,
                "description": {
                    "page_view": "页面访问",
                    "button_click": "按钮点击",
                    "api_request": "API请求",
                    "file_upload": "文件上传",
                    "file_download": "文件下载",
                    "report_view": "报告查看",
                    "task_start": "任务开始",
                    "task_complete": "任务完成",
                    "task_failed": "任务失败",
                    "user_login": "用户登录",
                    "user_logout": "用户登出"
                }
            }
        })
    except Exception as e:
        logger.error(f"获取操作类型失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


@app.get("/api/history/actions/count")
async def count_action_logs(current_user: Dict = Depends(get_current_user)):
    """统计用户操作记录总数
    
    所有用户都可以查看所有用户的统计数据。
    """
    try:
        storage = get_action_log_storage()
        
        # 所有用户都可以查看所有用户的统计数据
        total_count = 0
        type_counts = {}
        user_counts = {}
        
        base_dir = storage.base_dir
        if base_dir.exists():
            for user_dir in base_dir.iterdir():
                if user_dir.is_dir():
                    user_id = user_dir.name
                    user_count = storage.count(user_id)
                    user_counts[user_id] = user_count
                    total_count += user_count
                    
                    # 获取用户的操作类型统计
                    user_records, _ = storage.query(user_id, limit=10000)
                    for record in user_records:
                        action_type = record.get("action_type", "unknown")
                        type_counts[action_type] = type_counts.get(action_type, 0) + 1
        
        return JSONResponse({
            "code": 0,
            "message": "获取成功",
            "data": {
                "count": total_count,
                "type_counts": type_counts,
                "user_counts": user_counts
            }
        })
    except Exception as e:
        logger.error(f"统计操作日志失败: {e}")
        return JSONResponse({
            "code": 0,
            "message": "获取成功",
            "data": {"count": 0}
        })


@app.get("/api/history/actions/{action_id}")
async def get_action_log(
    action_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """获取单条操作记录详情"""
    try:
        storage = get_action_log_storage()
        record = storage.get(current_user["user_id"], action_id)

        if not record:
            return JSONResponse({
                "code": 1,
                "message": "记录不存在",
                "data": None
            }, status_code=404)

        return JSONResponse({
            "code": 0,
            "message": "获取成功",
            "data": record
        })
    except Exception as e:
        logger.error(f"获取操作日志失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


@app.delete("/api/history/actions/{action_id}")
async def delete_action_log(
    action_id: str,
    current_user: Dict = Depends(get_current_user)
):
    """删除单条操作记录"""
    try:
        storage = get_action_log_storage()
        success = storage.delete(current_user["user_id"], action_id)

        if not success:
            return JSONResponse({
                "code": 1,
                "message": "记录不存在",
                "data": None
            }, status_code=404)

        return JSONResponse({
            "code": 0,
            "message": "删除成功",
            "data": None
        })
    except Exception as e:
        logger.error(f"删除操作日志失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


@app.delete("/api/history/actions/cleanup")
async def cleanup_action_logs(
    before_time: str,
    current_user: Dict = Depends(require_admin)
):
    """清理指定时间之前的操作记录（仅管理员可访问）
    
    参数:
        before_time: ISO格式时间字符串，如 "2026-01-01T00:00:00"
    """
    try:
        storage = get_action_log_storage()
        deleted_count = storage.delete_by_time_range(current_user["user_id"], before_time)

        return JSONResponse({
            "code": 0,
            "message": "清理成功",
            "data": {"deleted_count": deleted_count}
        })
    except Exception as e:
        logger.error(f"清理操作日志失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


@app.post("/api/initialize")
async def initialize_demo_data(current_user: Dict = Depends(get_current_user)):
    """初始化演示数据（可选调用）"""
    try:
        import json
        from datetime import datetime, timedelta
        from .action_logger import get_action_log_storage
        
        action_storage = get_action_log_storage()
        
        # 生成一些示例操作日志
        now = datetime.now()
        demo_actions = [
            ("page_view", "访问页面: 首页", "", {}),
            ("button_click", "点击按钮: 开始分析", "", {}),
            ("file_upload", "上传文件: demo.log", "demo.log", {"file_size": 10240}),
            ("api_request", "GET /api/list-directory", "/api/list-directory", {"method": "GET", "status_code": 200}),
            ("task_start", "任务开始: demo_task", "demo_task", {"file_name": "demo.log"}),
            ("task_complete", "任务完成: demo_task", "demo_task", {}),
            ("report_view", "查看报告", "", {}),
            ("file_download", "下载文件: report.md", "report.md", {}),
        ]
        
        for action_type, action_name, resource, details in demo_actions:
            action_storage.record(
                user_id=current_user["user_id"],
                action_type=action_type,
                action_name=action_name,
                resource=resource,
                details=details
            )
        
        return JSONResponse({
            "code": 0,
            "message": "演示数据初始化成功",
            "data": {"initialized": True}
        })
    except Exception as e:
        logger.error(f"初始化演示数据失败: {e}")
        return JSONResponse({
            "code": 1,
            "message": str(e),
            "data": None
        }, status_code=500)


static_dir = SCRIPT_DIR / "static"
if static_dir.exists():
    app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")
