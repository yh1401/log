"""Report generator for creating structured analysis reports."""

import os
import json
import logging
import importlib
from dataclasses import dataclass, field
from datetime import datetime
from typing import List, Dict, Any, Optional

import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from ..processor.chunk_processor import ProcessingResult
from ..llm.client import AnalysisResult
from ..utils.helpers import ensure_dir, get_file_size_str
from .error_merger import ErrorMerger, MergeConfig


def _get_weasyprint():
    """延迟加载 WeasyPrint，避免启动时的依赖问题"""
    try:
        weasyprint = importlib.import_module('weasyprint')
        return weasyprint.HTML
    except ImportError:
        return None


def markdown_to_html(content: str) -> str:
    """
    使用 python-markdown 库将 Markdown 转换为 HTML。
    支持完整的 Markdown 语法，包括：
    - 标题：#、##、### 等
    - 列表：有序和无序列表
    - 粗体/斜体：**、*
    - 代码块和行内代码
    - 链接和图片
    - 表格
    - 引用块等
    
    Args:
        content: Markdown 格式的文本内容
        
    Returns:
        HTML 格式的文本内容
    """
    if not content:
        return ""
    
    # 使用 python-markdown 进行转换
    # extensions 参数添加额外功能支持
    return markdown.markdown(content, extensions=['tables', 'fenced_code', 'sane_lists'])


def _create_apple_style_table(doc, data, headers):
    """
    创建苹果风格的表格
    """
    rows = len(data) + 1  # +1 for header
    cols = len(headers)
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.autofit = False
    
    # 设置列宽
    for col in table.columns:
        col.width = Cm(7 / cols)  # 平均分配宽度
    
    # 设置表头样式
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'SF Pro Display'
                run.font.size = Pt(10)
                run.bold = True
                run.font.color.rgb = RGBColor(77, 77, 77)
                run.font.underline = True
    
    # 设置数据行样式
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]
        for col_idx, cell_value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.space_after = Pt(6)
                for run in paragraph.runs:
                    run.font.name = 'SF Pro Display'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(51, 51, 51)
    
    return table


def _format_json_to_table(doc, json_str, table_name):
    """
    将JSON格式的字符串转换为结构化表格
    """
    import json as json_module
    
    try:
        # 尝试解析JSON
        data = json_module.loads(json_str)
        
        if isinstance(data, list):
            # 如果是列表，提取公共字段作为表头
            if data and isinstance(data[0], dict):
                headers = list(data[0].keys())
                rows = []
                for item in data:
                    row = []
                    for header in headers:
                        value = item.get(header, '')
                        # 处理嵌套结构
                        if isinstance(value, dict):
                            value = json_module.dumps(value, ensure_ascii=False)[:50] + '...' if len(str(value)) > 50 else json_module.dumps(value, ensure_ascii=False)
                        elif isinstance(value, list):
                            value = str(value)[:50] + '...' if len(str(value)) > 50 else str(value)
                        row.append(str(value))
                    rows.append(row)
                
                # 添加表格名称
                p = doc.add_paragraph()
                run = p.add_run(f"**{table_name}**")
                run.font.name = 'SF Pro Display'
                run.font.size = Pt(11)
                run.bold = True
                
                # 创建表格
                _create_apple_style_table(doc, rows, headers)
                
                # 添加空行
                doc.add_paragraph()
                return True
        
        elif isinstance(data, dict):
            # 如果是字典，转换为两列表格
            rows = []
            for key, value in data.items():
                if isinstance(value, dict):
                    value = json_module.dumps(value, ensure_ascii=False)[:80] + '...' if len(str(value)) > 80 else json_module.dumps(value, ensure_ascii=False)
                elif isinstance(value, list):
                    value = str(value)[:80] + '...' if len(str(value)) > 80 else str(value)
                rows.append([key, str(value)])
            
            # 添加表格名称
            p = doc.add_paragraph()
            run = p.add_run(f"**{table_name}**")
            run.font.name = 'SF Pro Display'
            run.font.size = Pt(11)
            run.bold = True
            
            # 创建表格
            _create_apple_style_table(doc, rows, ['字段', '值'])
            
            # 添加空行
            doc.add_paragraph()
            return True
    
    except (json_module.JSONDecodeError, ValueError):
        pass
    
    return False


def markdown_to_docx(content: str, output_path: str) -> None:
    """
    将 Markdown 内容转换为 Word 文档（.docx）。
    采用苹果公司官网风格设计。
    
    支持的格式：
    - 标题：#、##、###、####
    - 无序列表：- 开头的行
    - 有序列表：1.、2. 等开头的行
    - 粗体：**text**
    - 斜体：*text*
    - 行内代码：`code`
    - 表格：Markdown 表格格式
    - 引用块：> 开头的行
    - JSON数据自动转换为表格
    - 换行符
    """
    doc = Document()
    
    # 设置默认样式（苹果风格）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SF Pro Display'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 51, 51)
    style.paragraph_format.line_spacing = 1.6
    style.paragraph_format.space_after = Pt(10)
    
    # 设置标题样式
    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_font = heading_style.font
        heading_font.name = 'SF Pro Display'
        heading_font.color.rgb = RGBColor(30, 30, 30)
        heading_font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(8)
    
    # 设置页面边距（苹果风格：上下2.54cm，左右2.54cm）
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 处理标题
        if line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
        elif line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(20, 20, 20)
            i += 1
        
        # 处理表格
        elif line.startswith('|') and '---' in line:
            # 找到表格的所有行
            table_lines = []
            # 先回退一行，可能是表头
            if i > 0 and lines[i-1].strip().startswith('|'):
                table_lines.append(lines[i-1].strip())
                i -= 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # 解析表格
            if len(table_lines) >= 2:
                # 获取列数
                col_count = len(table_lines[0].split('|')) - 1
                
                # 创建表格
                table = doc.add_table(rows=len(table_lines)-1, cols=col_count)
                table.style = 'Table Grid'
                table.autofit = False
                
                # 设置列宽
                for col in table.columns:
                    col.width = Cm(14 / col_count)
                
                # 填充表格内容
                for row_idx, table_line in enumerate(table_lines):
                    if row_idx == 1:
                        # 跳过分隔线行
                        continue
                    
                    cells = table_line.split('|')[1:-1]  # 去掉首尾的 |
                    adjusted_row_idx = row_idx - (1 if row_idx > 1 else 0)
                    
                    for col_idx, cell_text in enumerate(cells):
                        cell = table.cell(adjusted_row_idx, col_idx)
                        cell.text = cell_text.strip()
                        
                        # 设置单元格样式
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph.space_after = Pt(0)
                            for run in paragraph.runs:
                                run.font.name = 'SF Pro Display'
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(68, 68, 68)
                
                # 添加表格后空一行
                if i < len(lines) and lines[i].strip():
                    doc.add_paragraph()
        
        # 处理JSON数据（检测并转换为表格）
        elif (line.startswith('{') and '}' in line) or (line.startswith('[') and ']' in line):
            # 收集多行JSON数据
            json_lines = [line]
            brace_count = line.count('{') - line.count('}') + line.count('[') - line.count(']')
            
            while i + 1 < len(lines) and brace_count > 0:
                i += 1
                next_line = lines[i].strip()
                json_lines.append(next_line)
                brace_count += next_line.count('{') - next_line.count('}') + next_line.count('[') - next_line.count(']')
            
            full_json = ' '.join(json_lines)
            
            # 尝试转换为表格
            if not _format_json_to_table(doc, full_json, "数据详情"):
                # 如果转换失败，作为普通文本处理
                p = doc.add_paragraph()
                run = p.add_run(full_json)
                run.font.name = 'SF Mono'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)
            
            i += 1
            continue
        
        # 处理无序列表
        elif line.startswith('- '):
            # 找到连续的列表项
            list_items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                list_items.append(lines[i].strip()[2:])
                i += 1
            
            # 添加列表
            for item in list_items:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.add_run('• ').bold = True
                # 处理粗体、斜体和代码
                parts = parse_markdown_inline(item)
                for part in parts:
                    run = p.add_run(part['text'])
                    if part.get('bold'):
                        run.bold = True
                    if part.get('italic'):
                        run.italic = True
                    if part.get('code'):
                        run.font.name = 'SF Mono'
                        run.font.color.rgb = RGBColor(0, 122, 255)
        
        # 处理有序列表
        elif line.replace('.', '').strip().isdigit() and line.count('.') == 1:
            # 尝试解析有序列表
            import re
            match = re.match(r'^(\d+)\.\s+(.+)$', line)
            if match:
                # 找到连续的列表项
                list_items = []
                expected_num = int(match.group(1))
                while i < len(lines):
                    line_i = lines[i].strip()
                    match_i = re.match(r'^(\d+)\.\s+(.+)$', line_i)
                    if match_i and int(match_i.group(1)) == expected_num:
                        list_items.append(match_i.group(2))
                        expected_num += 1
                        i += 1
                    else:
                        break
                
                # 添加有序列表
                for idx, item in enumerate(list_items):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.add_run(f'{idx + 1}. ').bold = True
                    # 处理粗体、斜体和代码
                    parts = parse_markdown_inline(item)
                    for part in parts:
                        run = p.add_run(part['text'])
                        if part.get('bold'):
                            run.bold = True
                        if part.get('italic'):
                            run.italic = True
                        if part.get('code'):
                            run.font.name = 'SF Mono'
                            run.font.color.rgb = RGBColor(0, 122, 255)
                continue
        
        # 处理引用块
        elif line.startswith('> '):
            # 找到连续的引用行
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            
            # 添加引用段落
            p = doc.add_paragraph()
            p.style = 'Quote'
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.8)
            p.paragraph_format.border_left.width = Pt(3)
            p.paragraph_format.border_left.color = RGBColor(0, 122, 255)
            run = p.add_run('\n'.join(quote_lines))
            run.font.name = 'SF Pro Display'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(120, 120, 120)
            run.italic = True
        
        # 处理空行
        elif not line:
            i += 1
        
        # 处理普通文本段落
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.75)
            # 处理粗体、斜体和代码
            parts = parse_markdown_inline(line)
            for part in parts:
                run = p.add_run(part['text'])
                if part.get('bold'):
                    run.bold = True
                if part.get('italic'):
                    run.italic = True
                if part.get('code'):
                    run.font.name = 'SF Mono'
                    run.font.color.rgb = RGBColor(0, 122, 255)
            i += 1
    
    doc.save(output_path)


def markdown_to_pdf(content: str, output_path: str) -> None:
    """
    将 Markdown 内容转换为 PDF 文档。
    支持完整的 Markdown 语法，包括表格、代码块等。
    
    Args:
        content: Markdown 格式的文本内容
        output_path: 输出 PDF 文档的路径
    
    Raises:
        ImportError: 如果 WeasyPrint 不可用
    """
    # 延迟加载 WeasyPrint
    WeasyHTML = _get_weasyprint()
    
    if not WeasyHTML:
        raise ImportError(
            "WeasyPrint 不可用，请安装系统依赖。\n"
            "macOS: brew install pygobject3 gtk+3 libffi\n"
            "Ubuntu/Debian: apt-get install libgirepository1.0-dev libcairo2-dev libpango1.0-dev\n"
            "Windows: 需要安装 GTK+ 运行时环境"
        )
    
    # 先将 Markdown 转换为 HTML
    html_content = markdown_to_html(content)
    
    # 添加基本样式
    full_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@400;600;700&display=swap');
        
        body {{
            font-family: 'Noto Serif SC', 'Songti SC', serif;
            font-size: 11pt;
            line-height: 1.6;
            margin: 2cm;
            color: #333;
        }}
        
        h1 {{
            text-align: center;
            font-size: 18pt;
            font-weight: 700;
            margin-bottom: 1cm;
            color: #1a1a1a;
        }}
        
        h2 {{
            font-size: 14pt;
            font-weight: 600;
            margin-top: 1cm;
            margin-bottom: 0.5cm;
            border-bottom: 2px solid #2c3e50;
            padding-bottom: 5px;
            color: #2c3e50;
        }}
        
        h3 {{
            font-size: 12pt;
            font-weight: 600;
            margin-top: 0.8cm;
            margin-bottom: 0.3cm;
            color: #34495e;
        }}
        
        h4 {{
            font-size: 11pt;
            font-weight: 600;
            margin-top: 0.6cm;
            margin-bottom: 0.2cm;
            color: #7f8c8d;
        }}
        
        p {{
            margin: 0.3cm 0;
            text-align: justify;
        }}
        
        ul, ol {{
            margin: 0.3cm 0;
            padding-left: 1.5cm;
        }}
        
        li {{
            margin: 0.2cm 0;
        }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 0.5cm 0;
            font-size: 10pt;
        }}
        
        th, td {{
            border: 1px solid #ddd;
            padding: 8px 12px;
            text-align: center;
        }}
        
        th {{
            background-color: #f8f9fa;
            font-weight: 600;
        }}
        
        tr:nth-child(even) {{
            background-color: #fafafa;
        }}
        
        code {{
            font-family: 'SF Mono', 'Consolas', monospace;
            font-size: 9pt;
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            color: #e74c3c;
        }}
        
        pre {{
            background-color: #2d2d2d;
            color: #f8f8f2;
            padding: 12px;
            border-radius: 6px;
            overflow-x: auto;
            font-family: 'SF Mono', 'Consolas', monospace;
            font-size: 9pt;
            margin: 0.5cm 0;
        }}
        
        pre code {{
            background: none;
            color: inherit;
            padding: 0;
        }}
        
        blockquote {{
            border-left: 4px solid #3498db;
            margin: 0.5cm 0;
            padding-left: 15px;
            color: #7f8c8d;
            font-style: italic;
        }}
        
        strong {{
            font-weight: 600;
        }}
        
        em {{
            font-style: italic;
        }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>
    """
    
    # 使用 WeasyPrint 生成 PDF
    WeasyHTML(string=full_html).write_pdf(output_path)


def parse_markdown_inline(text: str) -> List[Dict[str, Any]]:
    """
    解析行内 Markdown 格式（粗体、斜体和代码）。
    
    Args:
        text: 包含行内 Markdown 的文本
        
    Returns:
        解析后的片段列表，每个片段包含 text 和格式标记
    """
    parts = []
    i = 0
    
    while i < len(text):
        # 检测粗体 **...**（必须在斜体之前检测，避免冲突）
        if i + 1 < len(text) and text[i:i+2] == '**':
            end = text.find('**', i + 2)
            if end != -1:
                parts.append({'text': text[i+2:end], 'bold': True})
                i = end + 2
                continue
        
        # 检测斜体 *...*（注意：不是粗体的一部分）
        if text[i] == '*' and (i == 0 or text[i-1] != '*'):
            end = text.find('*', i + 1)
            # 确保结束的 * 不是下一个粗体的开始
            if end != -1 and (end + 1 >= len(text) or text[end+1] != '*'):
                parts.append({'text': text[i+1:end], 'italic': True})
                i = end + 1
                continue
        
        # 检测代码 `...`
        if text[i] == '`':
            end = text.find('`', i + 1)
            if end != -1:
                parts.append({'text': text[i+1:end], 'code': True})
                i = end + 1
                continue
        
        # 普通文本
        parts.append({'text': text[i]})
        i += 1
    
    # 合并相邻的普通文本
    result = []
    for part in parts:
        is_special = part.get('bold') or part.get('italic') or part.get('code')
        if result and 'text' in part and not is_special:
            result[-1]['text'] += part['text']
        else:
            result.append(part)
    
    return result


@dataclass
class ReportSection:
    title: str
    content: str
    section_type: str = "text"
    data: Optional[Dict[str, Any]] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            'title': self.title,
            'content': self.content,
            'section_type': self.section_type,
            'data': self.data
        }


@dataclass
class Report:
    title: str
    generated_at: datetime
    file_path: str
    file_size: str
    total_lines: int
    total_errors: int
    total_warnings: int
    sections: List[ReportSection] = field(default_factory=list)
    summary: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            'title': self.title,
            'generated_at': self.generated_at.isoformat(),
            'file_path': self.file_path,
            'file_size': self.file_size,
            'total_lines': self.total_lines,
            'total_errors': self.total_errors,
            'total_warnings': self.total_warnings,
            'summary': self.summary,
            'sections': [s.to_dict() for s in self.sections],
            'metadata': self.metadata
        }

    def to_markdown(self) -> str:
        lines = [
            f"# {self.title}\n",
            f"**生成时间**: {self.generated_at.strftime('%Y-%m-%d %H:%M:%S')}\n",
            f"**文件**: {self.file_path}\n",
            f"**文件大小**: {self.file_size}\n",
            f"**总行数**: {self.total_lines:,}\n",
            f"**总错误数**: {self.total_errors:,}\n",
            f"**总警告数**: {self.total_warnings:,}\n",
            "---\n"
        ]

        # 收集所有section的数据用于报告末尾的汇总
        all_data = {}

        for section in self.sections:
            lines.append(f"## {section.title}\n")
            lines.append(f"{section.content}\n")
            lines.append("\n")
            
            # 收集数据用于汇总
            if section.data:
                all_data[section.section_type] = section.data

        if self.summary:
            lines.append("---\n\n")
            lines.append(f"## 总体摘要\n\n{self.summary}\n")

        return "".join(lines)


class ReportGenerator:
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        ensure_dir(output_dir)

    def generate_combined_report(self, results: List[ProcessingResult]) -> Report:
        """生成多个文件的综合分析报告"""
        total_lines = sum(r.total_lines for r in results)
        total_errors = sum(
            r.statistics.get('by_level', {}).get('ERROR', 0) +
            r.statistics.get('by_level', {}).get('FATAL', 0)
            for r in results
        )
        total_warnings = sum(
            r.statistics.get('by_level', {}).get('WARN', 0)
            for r in results
        )
        processed_files = sum(1 for r in results if r.status == 'completed')

        report = Report(
            title=f"综合日志分析报告 - {len(results)} 个文件",
            generated_at=datetime.now(),
            file_path=f"{processed_files} 个文件",
            file_size=f"{len(results)} 个文件",
            total_lines=total_lines,
            total_errors=total_errors,
            total_warnings=total_warnings,
            metadata={
                'total_files': len(results),
                'processed_files': processed_files,
                'file_names': [os.path.basename(r.file_path) for r in results]
            }
        )

        # 添加概览
        report.sections.append(self._create_combined_overview_section(results))
        
        # 添加各文件摘要
        report.sections.append(self._create_file_summaries_section(results))
        
        # 添加综合统计
        report.sections.append(self._create_combined_statistics_section(results))
        
        # 添加综合建议
        report.sections.append(self._create_combined_suggestions_section(results))

        report.summary = self._generate_combined_summary(results)

        return report

    def generate_report(self, result: ProcessingResult) -> Report:
        total_errors = (
            result.statistics.get('by_level', {}).get('ERROR', 0) +
            result.statistics.get('by_level', {}).get('FATAL', 0)
        )
        total_warnings = result.statistics.get('by_level', {}).get('WARN', 0)

        report = Report(
            title=f"日志分析报告 - {os.path.basename(result.file_path)}",
            generated_at=datetime.now(),
            file_path=result.file_path,
            file_size=get_file_size_str(result.file_path),
            total_lines=result.total_lines,
            total_errors=total_errors,
            total_warnings=total_warnings,
            metadata={
                'processed_lines': result.processed_lines,
                'total_chunks': result.total_chunks,
                'completed_chunks': result.completed_chunks,
                'status': result.status
            }
        )

        report.sections.append(self._create_overview_section(result))
        report.sections.append(self._create_statistics_section(result))
        report.sections.append(self._create_combined_error_section(result))
        report.sections.append(self._create_trends_timeline_section(result))
        report.sections.append(self._create_root_cause_analysis_section(result))
        report.sections.append(self._create_disposal_remediation_section(result))

        report.summary = self._generate_summary(result)

        return report

    def _create_overview_section(self, result: ProcessingResult) -> ReportSection:
        overview_data = {
            '处理状态': result.status,
            '文件路径': result.file_path,
            '总行数': result.total_lines,
            '已处理行数': result.processed_lines,
            '处理进度': f"{result.get_progress_percentage():.2f}%",
            '总块数': result.total_chunks,
            '已完成块数': result.completed_chunks
        }

        content = f"""
**处理状态**: {result.status}
**处理进度**: {result.get_progress_percentage():.2f}%
**已处理**: {result.processed_lines:,} / {result.total_lines:,} 行
**已完成**: {result.completed_chunks} / {result.total_chunks} 块
"""
        if result.error_message:
            content += f"\n**错误信息**: {result.error_message}\n"

        return ReportSection(
            title="1. 处理概览",
            content=content.strip(),
            section_type="overview",
            data=overview_data
        )

    def _create_statistics_section(self, result: ProcessingResult) -> ReportSection:
        by_level = result.statistics.get('by_level', {})
        error_types = result.statistics.get('error_types', {})
        top_classes = result.statistics.get('top_classes', {})

        stats_data = {
            '错误级别分布': by_level,
            '错误类型统计': dict(list(error_types.items())[:20]),
            '高频错误类': dict(list(top_classes.items())[:20])
        }

        level_lines = [f"- **{level}**: {count:,}" for level, count in sorted(by_level.items(), key=lambda x: x[1], reverse=True)]

        content = f"""
### 错误级别分布
{chr(10).join(level_lines)}

### 错误类型统计 (Top 20)
"""
        for error_type, count in list(error_types.items())[:20]:
            content += f"- {error_type}: {count:,}\n"

        content += "\n### 高频错误类 (Top 20)\n"
        for class_name, count in list(top_classes.items())[:20]:
            content += f"- {class_name}: {count:,}\n"

        return ReportSection(
            title="2. 统计分析",
            content=content.strip(),
            section_type="statistics",
            data=stats_data
        )

    def _create_error_analysis_section(self, result: ProcessingResult) -> ReportSection:
        # 创建智能错误合并器，使用可配置的合并策略
        merge_config = MergeConfig(
            semantic_similarity_threshold=0.75,
            max_examples_per_group=3,
            max_groups=10,
            enable_semantic_merging=True,
            merge_by_error_type=True,
            merge_by_message_pattern=True
        )
        error_merger = ErrorMerger(merge_config)
        
        # 使用智能合并功能合并错误
        merged_errors = error_merger.merge_from_analysis_results(result.analysis_results)

        errors_data = {'关键错误': merged_errors}

        content = "### 关键错误分析\n\n"
        for idx, error in enumerate(merged_errors, 1):
            error_type = error.get('error_type', 'Unknown')
            description = error.get('description', '')
            count = error.get('count', 0)
            severity = error.get('severity', 'medium')
            
            content += f"#### {idx}. {error_type}\n"
            content += f"- **描述**: {description}\n"
            content += f"- **出现次数**: {count:,}\n"
            content += f"- **严重程度**: {severity}\n"
            
            # 如果有影响的类，显示出来
            affected_classes = error.get('affected_classes', [])
            if affected_classes:
                content += f"- **影响类**: {', '.join(affected_classes[:5])}"
                if len(affected_classes) > 5:
                    content += f" 等{len(affected_classes)}个类"
                content += "\n"
            
            # 如果有示例，显示示例消息
            examples = error.get('examples', [])
            if examples:
                content += f"- **示例消息**:\n"
                for i, example in enumerate(examples[:3], 1):
                    msg = example.get('message', '')[:80] + "..." if len(example.get('message', '')) > 80 else example.get('message', '')
                    content += f"  {i}. {msg}\n"
            
            content += "\n"

        return ReportSection(
            title="关键错误分析",
            content=content.strip(),
            section_type="error_analysis",
            data=errors_data
        )

    def _create_pattern_analysis_section(self, result: ProcessingResult) -> ReportSection:
        all_patterns: Dict[str, int] = {}
        for analysis in result.analysis_results:
            freq_stats = analysis.frequency_stats
            if isinstance(freq_stats, dict):
                for pattern_type, value in freq_stats.items():
                    if isinstance(value, int):
                        all_patterns[pattern_type] = all_patterns.get(pattern_type, 0) + value
                    elif isinstance(value, dict):
                        for sub_key, sub_value in value.items():
                            if isinstance(sub_value, int):
                                combined_key = f"{pattern_type}.{sub_key}"
                                all_patterns[combined_key] = all_patterns.get(combined_key, 0) + sub_value

        patterns_sorted = sorted(all_patterns.items(), key=lambda x: x[1], reverse=True)
        patterns_data = {'错误模式统计': dict(patterns_sorted)}

        content = "### 错误模式识别\n\n"
        for pattern_type, count in patterns_sorted:
            content += f"- **{pattern_type}**: {count:,} 次\n"

        return ReportSection(
            title="错误模式分析",
            content=content.strip(),
            section_type="pattern_analysis",
            data=patterns_data
        )

    def _create_trends_section(self, result: ProcessingResult) -> ReportSection:
        all_trends = []
        for analysis in result.analysis_results:
            all_trends.extend(analysis.trends)

        seen = set()
        unique_trends = []
        for trend in all_trends:
            if trend not in seen:
                seen.add(trend)
                unique_trends.append(trend)

        trends_data = {'趋势识别': unique_trends}

        content = "### 趋势识别\n\n"
        content += "通过对日志数据的深入分析，识别出以下关键趋势和模式：\n\n"
        content += "**识别到的趋势：**\n\n"
        for idx, trend in enumerate(unique_trends[:10], 1):
            content += f"{idx}. **{trend}**\n"

        return ReportSection(
            title="趋势识别",
            content=content.strip(),
            section_type="trends",
            data=trends_data
        )

    def _create_suggestions_section(self, result: ProcessingResult) -> ReportSection:
        all_ops_suggestions = []
        all_dev_suggestions = []
        
        for analysis in result.analysis_results:
            # 收集运维建议
            for suggestion in analysis.ops_suggestions:
                if suggestion and suggestion not in all_ops_suggestions:
                    all_ops_suggestions.append(suggestion)
            
            # 收集开发建议
            for suggestion in analysis.dev_suggestions:
                if suggestion and suggestion not in all_dev_suggestions:
                    all_dev_suggestions.append(suggestion)
        
        # 处理建议数据，提取category和suggestion字段
        processed_ops = []
        for s in all_ops_suggestions[:3]:
            if isinstance(s, dict):
                processed_ops.append({
                    'category': s.get('category', ''),
                    'suggestion': s.get('suggestion', '')
                })
            else:
                processed_ops.append({'category': '', 'suggestion': str(s)})
        
        processed_dev = []
        for s in all_dev_suggestions[:3]:
            if isinstance(s, dict):
                processed_dev.append({
                    'category': s.get('category', ''),
                    'suggestion': s.get('suggestion', '')
                })
            else:
                processed_dev.append({'category': '', 'suggestion': str(s)})
        
        suggestions_data = {
            '运维建议': processed_ops,
            '开发建议': processed_dev
        }
        
        content = "### 解决建议\n\n"
        
        # 运维建议部分
        if processed_ops:
            content += "#### 运维建议\n\n"
            for idx, item in enumerate(processed_ops, 1):
                category = item.get('category', '')
                suggestion = item.get('suggestion', '')
                if category:
                    content += f"{idx}. **{category}**: {suggestion}\n"
                else:
                    content += f"{idx}. {suggestion}\n"
            content += "\n"
        
        # 开发建议部分
        if processed_dev:
            content += "#### 开发建议\n\n"
            for idx, item in enumerate(processed_dev, 1):
                category = item.get('category', '')
                suggestion = item.get('suggestion', '')
                if category:
                    content += f"{idx}. **{category}**: {suggestion}\n"
                else:
                    content += f"{idx}. {suggestion}\n"
        
        return ReportSection(
            title="解决建议",
            content=content.strip(),
            section_type="suggestions",
            data=suggestions_data
        )

    def _create_timeline_section(self, result: ProcessingResult) -> ReportSection:
        all_timelines = []
        for analysis in result.analysis_results:
            if analysis.timeline:
                all_timelines.append(analysis.timeline)

        if not all_timelines:
            return ReportSection(
                title="故障时间线",
                content="暂无故障时间线数据",
                section_type="timeline",
                data={}
            )

        timeline_data = all_timelines[0] if all_timelines else {}

        description = timeline_data.get('description', '')
        key_events = timeline_data.get('key_events', [])
        total_duration = timeline_data.get('total_duration', '')

        content = f"### 故障时间线\n\n"
        content += "根据日志分析，故障事件的时间线如下：\n\n"
        
        if description:
            content += f"**事件概述**: {description}\n\n"
        if total_duration:
            content += f"**持续时长**: {total_duration}\n\n"

        if key_events:
            content += "#### 📊 关键事件序列\n\n"
            content += "| 序号 | 时间 | 事件类型 | 描述 |\n"
            content += "|------|------|----------|------|\n"
            
            event_type_names = {
                'first_abnormal': '🔴 首次异常',
                'peak_error': '🔥 错误峰值',
                'recovery': '🟢 恢复',
                'fault_confirmed': '⚠️ 故障确认'
            }
            
            for idx, event in enumerate(key_events, 1):
                event_time = event.get('time', 'N/A')
                event_type = event.get('event_type', 'unknown')
                event_desc = event.get('description', '')

                display_type = event_type_names.get(event_type, event_type)
                content += f"| {idx} | {event_time} | {display_type} | {event_desc} |\n"

            content += "\n"

        return ReportSection(
            title="一、故障时间线（Fault Timeline）",
            content=content.strip(),
            section_type="timeline",
            data={'timeline': timeline_data}
        )

    def _create_root_cause_section(self, result: ProcessingResult) -> ReportSection:
        all_root_causes = []
        for analysis in result.analysis_results:
            if analysis.root_cause:
                all_root_causes.append(analysis.root_cause)

        if not all_root_causes:
            return ReportSection(
                title="根因推断",
                content="暂无根因分析数据",
                section_type="root_cause",
                data={}
            )

        root_cause_data = all_root_causes[0] if all_root_causes else {}

        direct_cause = root_cause_data.get('direct_cause', '')
        fundamental_cause = root_cause_data.get('fundamental_cause', '')
        confidence = root_cause_data.get('confidence', '')
        reasoning = root_cause_data.get('reasoning', '')

        content = "### 根因推断（Root Cause Inference）\n\n"
        content += f"**直接原因（Direct Cause）**: {direct_cause}\n\n"
        content += f"**根本原因（Root Cause）**: {fundamental_cause}\n\n"

        if confidence:
            confidence_display = {'high': '高', 'medium': '中', 'low': '低'}.get(confidence.lower(), confidence)
            content += f"**置信度**: {confidence_display}\n\n"
        if reasoning:
            content += f"**推断依据**: {reasoning}\n\n"

        return ReportSection(
            title="二、根因推断（Root Cause Inference）",
            content=content.strip(),
            section_type="root_cause",
            data={'root_cause': root_cause_data}
        )

    def _create_causal_chain_section(self, result: ProcessingResult) -> ReportSection:
        all_causal_chains = []
        for analysis in result.analysis_results:
            if analysis.causal_chain:
                all_causal_chains.append(analysis.causal_chain)

        if not all_causal_chains:
            return ReportSection(
                title="故障因果链",
                content="暂无因果链数据",
                section_type="causal_chain",
                data={}
            )

        causal_chain_data = all_causal_chains[0] if all_causal_chains else {}

        chain_description = causal_chain_data.get('chain_description', '')
        chain_steps = causal_chain_data.get('chain_steps', [])

        content = "### 故障因果链（Causal Chain）\n\n"
        if chain_description:
            content += f"{chain_description}\n\n"

        if chain_steps:
            content += "#### 因果传播路径\n\n"
            for step in chain_steps:
                step_num = step.get('step', 0)
                cause = step.get('cause', '')
                effect = step.get('effect', '')
                evidence = step.get('evidence', '')
                timestamp = step.get('timestamp', '')

                content += f"**步骤 {step_num}**\n"
                if timestamp:
                    content += f"- 时间: {timestamp}\n"
                content += f"- 原因: {cause}\n"
                content += f"- 结果: {effect}\n"
                if evidence:
                    content += f"- 证据: {evidence}\n"
                content += "\n"

        return ReportSection(
            title="三、故障因果链（Causal Chain）",
            content=content.strip(),
            section_type="causal_chain",
            data={'causal_chain': causal_chain_data}
        )

    def _create_evidence_chain_section(self, result: ProcessingResult) -> ReportSection:
        all_evidence_chains = []
        for analysis in result.analysis_results:
            if analysis.evidence_chain:
                all_evidence_chains.append(analysis.evidence_chain)

        if not all_evidence_chains:
            return ReportSection(
                title="证据链",
                content="暂无证据链数据",
                section_type="evidence_chain",
                data={}
            )

        evidence_chain_data = all_evidence_chains[0] if all_evidence_chains else {}

        description = evidence_chain_data.get('description', '')
        evidences = evidence_chain_data.get('evidences', [])

        content = "### 证据链（Evidence Chain）\n\n"
        if description:
            content += f"{description}\n\n"

        if evidences:
            content += "#### 关键证据\n\n"
            relevance_names = {
                'direct': '直接关联',
                'indirect': '间接关联',
                'supporting': '辅助支撑'
            }
            evidence_type_names = {
                'log': '日志条目',
                'exception': '异常信息',
                'metric': '系统指标',
                'trace': '链路追踪'
            }
            for idx, ev in enumerate(evidences, 1):
                ev_timestamp = ev.get('timestamp', 'N/A')
                ev_type = ev.get('evidence_type', 'unknown')
                ev_content = ev.get('content', '')
                ev_relevance = ev.get('relevance', '')

                display_type = evidence_type_names.get(ev_type, ev_type)
                display_relevance = relevance_names.get(ev_relevance, ev_relevance)

                content += f"**证据 {idx}** [{ev_timestamp}]\n"
                content += f"- 类型: {display_type}\n"
                content += f"- 内容: {ev_content}\n"
                content += f"- 关联度: {display_relevance}\n\n"

        return ReportSection(
            title="四、证据链（Evidence Chain）",
            content=content.strip(),
            section_type="evidence_chain",
            data={'evidence_chain': evidence_chain_data}
        )

    def _create_response_actions_section(self, result: ProcessingResult) -> ReportSection:
        all_response_actions = []
        for analysis in result.analysis_results:
            if analysis.response_actions:
                all_response_actions.append(analysis.response_actions)

        if not all_response_actions:
            return ReportSection(
                title="处置动作建议",
                content="暂无处置动作建议数据",
                section_type="response_actions",
                data={}
            )

        response_data = all_response_actions[0] if all_response_actions else {}

        description = response_data.get('description', '')
        emergency_actions = response_data.get('emergency_actions', [])
        troubleshooting_actions = response_data.get('troubleshooting_actions', [])
        recovery_actions = response_data.get('recovery_actions', [])

        content = "### 处置动作建议（Immediate Response Actions）\n\n"
        if description:
            content += f"{description}\n\n"

        if emergency_actions:
            content += "#### 应急止血动作\n\n"
            for action in emergency_actions:
                action_name = action.get('action_name', '')
                timing = action.get('timing', '')
                steps = action.get('steps', '')
                expected_effect = action.get('expected_effect', '')
                notes = action.get('notes', '')

                content += f"- **{action_name}**\n"
                if timing:
                    content += f"  - 执行时机: {timing}\n"
                if steps:
                    content += f"  - 执行步骤: {steps}\n"
                if expected_effect:
                    content += f"  - 预期效果: {expected_effect}\n"
                if notes:
                    content += f"  - 注意事项: {notes}\n"
                content += "\n"

        if troubleshooting_actions:
            content += "#### 排查定位动作\n\n"
            for action in troubleshooting_actions:
                action_name = action.get('action_name', '')
                timing = action.get('timing', '')
                steps = action.get('steps', '')
                expected_effect = action.get('expected_effect', '')
                notes = action.get('notes', '')

                content += f"- **{action_name}**\n"
                if timing:
                    content += f"  - 执行时机: {timing}\n"
                if steps:
                    content += f"  - 执行步骤: {steps}\n"
                if expected_effect:
                    content += f"  - 预期效果: {expected_effect}\n"
                if notes:
                    content += f"  - 注意事项: {notes}\n"
                content += "\n"

        if recovery_actions:
            content += "#### 恢复验证动作\n\n"
            for action in recovery_actions:
                action_name = action.get('action_name', '')
                timing = action.get('timing', '')
                steps = action.get('steps', '')
                expected_effect = action.get('expected_effect', '')
                notes = action.get('notes', '')

                content += f"- **{action_name}**\n"
                if timing:
                    content += f"  - 执行时机: {timing}\n"
                if steps:
                    content += f"  - 执行步骤: {steps}\n"
                if expected_effect:
                    content += f"  - 预期效果: {expected_effect}\n"
                if notes:
                    content += f"  - 注意事项: {notes}\n"
                content += "\n"

        return ReportSection(
            title="五、处置动作建议（Immediate Response Actions）",
            content=content.strip(),
            section_type="response_actions",
            data={'response_actions': response_data}
        )

    def _create_remediation_section(self, result: ProcessingResult) -> ReportSection:
        all_remediations = []
        for analysis in result.analysis_results:
            if analysis.remediation:
                all_remediations.append(analysis.remediation)

        if not all_remediations:
            return ReportSection(
                title="整改建议",
                content="暂无整改建议数据",
                section_type="remediation",
                data={}
            )

        remediation_data = all_remediations[0] if all_remediations else {}

        immediate = remediation_data.get('immediate', [])
        root_cause_fix = remediation_data.get('root_cause_fix', [])
        architecture_monitoring = remediation_data.get('architecture_monitoring', [])

        content = "### 整改建议（Rectification Suggestions）\n\n"

        if immediate:
            content += "#### 立即处置（Immediate Actions）\n\n"
            content += "**目标**: 快速恢复服务，最小化业务影响 | **时间要求**: 1小时内可执行\n\n"
            for action in immediate:
                action_text = action.get('action', '')
                target = action.get('target', '')
                expected_effect = action.get('expected_effect', '')
                effort = action.get('effort_estimate', '')

                content += f"- **{action_text}**\n"
                if target:
                    content += f"  - 目标: {target}\n"
                if expected_effect:
                    content += f"  - 预期效果: {expected_effect}\n"
                if effort:
                    content += f"  - 工作量: {effort}\n"
                content += "\n"

        if root_cause_fix:
            content += "#### 根因解决（Root Cause Fix）\n\n"
            content += "**目标**: 彻底修复导致故障的根本问题 | **时间要求**: 短期Sprint内完成（1-2周）\n\n"
            for action in root_cause_fix:
                action_text = action.get('action', '')
                target = action.get('target', '')
                expected_effect = action.get('expected_effect', '')
                effort = action.get('effort_estimate', '')

                content += f"- **{action_text}**\n"
                if target:
                    content += f"  - 目标: {target}\n"
                if expected_effect:
                    content += f"  - 预期效果: {expected_effect}\n"
                if effort:
                    content += f"  - 工作量: {effort}\n"
                content += "\n"

        if architecture_monitoring:
            content += "#### 架构/监控改进（Architecture & Monitoring）\n\n"
            content += "**目标**: 提升系统整体稳定性和可观测性 | **时间要求**: 季度规划级别（1-3个月）\n\n"
            for action in architecture_monitoring:
                action_text = action.get('action', '')
                target = action.get('target', '')
                expected_effect = action.get('expected_effect', '')
                effort = action.get('effort_estimate', '')

                content += f"- **{action_text}**\n"
                if target:
                    content += f"  - 目标: {target}\n"
                if expected_effect:
                    content += f"  - 预期效果: {expected_effect}\n"
                if effort:
                    content += f"  - 工作量: {effort}\n"
                content += "\n"

        return ReportSection(
            title="六、整改建议（Rectification Suggestions）",
            content=content.strip(),
            section_type="remediation",
            data={'remediation': remediation_data}
        )

    def _create_combined_error_section(self, result: ProcessingResult) -> ReportSection:
        error_section = self._create_error_analysis_section(result)
        pattern_section = self._create_pattern_analysis_section(result)

        content = "### 错误分析\n\n"
        content += "#### 关键错误分析\n\n"
        content += error_section.content.replace("### 关键错误分析\n\n", "") + "\n\n"
        content += "#### 错误模式分析\n\n"
        content += pattern_section.content.replace("### 错误模式分析\n\n", "")

        combined_data = {
            '关键错误': error_section.data.get('关键错误', []),
            '错误模式统计': pattern_section.data.get('错误模式统计', {})
        }

        return ReportSection(
            title="3. 错误分析",
            content=content.strip(),
            section_type="combined_error",
            data=combined_data
        )

    def _create_trends_timeline_section(self, result: ProcessingResult) -> ReportSection:
        trends_section = self._create_trends_section(result)
        timeline_section = self._create_timeline_section(result)

        content = "### 趋势识别与故障时间线\n\n"
        content += "#### 趋势识别\n\n"
        content += trends_section.content.replace("### 趋势识别\n\n", "") + "\n\n"
        content += "#### 故障时间线\n\n"
        content += timeline_section.content.replace("### 故障时间线\n\n", "")

        combined_data = {
            '趋势识别': trends_section.data.get('趋势识别', []),
            '时间线': timeline_section.data.get('timeline', {})
        }

        return ReportSection(
            title="4. 趋势识别与故障时间线",
            content=content.strip(),
            section_type="trends_timeline",
            data=combined_data
        )

    def _create_root_cause_analysis_section(self, result: ProcessingResult) -> ReportSection:
        root_cause_section = self._create_root_cause_section(result)
        causal_chain_section = self._create_causal_chain_section(result)
        evidence_chain_section = self._create_evidence_chain_section(result)

        content = "### 根因分析\n\n"
        content += "#### 5.1 根因推断\n\n"
        content += root_cause_section.content.replace("### 根因推断（Root Cause Inference）\n\n", "") + "\n\n"
        content += "#### 5.2 故障因果链\n\n"
        content += causal_chain_section.content.replace("### 故障因果链（Causal Chain）\n\n", "") + "\n\n"
        content += "#### 5.3 证据链\n\n"
        content += evidence_chain_section.content.replace("### 证据链（Evidence Chain）\n\n", "")

        combined_data = {
            '根因推断': root_cause_section.data.get('root_cause', {}),
            '因果链': causal_chain_section.data.get('causal_chain', {}),
            '证据链': evidence_chain_section.data.get('evidence_chain', {})
        }

        return ReportSection(
            title="5. 根因分析",
            content=content.strip(),
            section_type="root_cause_analysis",
            data=combined_data
        )

    def _create_disposal_remediation_section(self, result: ProcessingResult) -> ReportSection:
        response_section = self._create_response_actions_section(result)
        remediation_section = self._create_remediation_section(result)
        suggestions_section = self._create_suggestions_section(result)

        content = "### 处置与整改建议\n\n"
        content += "#### 6.1 处置动作建议\n\n"
        content += response_section.content.replace("### 处置动作建议（Immediate Response Actions）\n\n", "") + "\n\n"
        content += "#### 6.2 整改建议\n\n"
        content += remediation_section.content.replace("### 整改建议（Rectification Suggestions）\n\n", "") + "\n\n"
        content += "#### 6.3 解决建议（运维+开发建议）\n\n"
        content += suggestions_section.content.replace("### 解决建议\n\n", "")

        combined_data = {
            '处置动作建议': response_section.data.get('response_actions', {}),
            '整改建议': remediation_section.data.get('remediation', {}),
            '解决建议': suggestions_section.data
        }

        return ReportSection(
            title="6. 处置与整改建议",
            content=content.strip(),
            section_type="disposal_remediation",
            data=combined_data
        )

    def _generate_summary(self, result: ProcessingResult) -> str:
        summaries = []
        for analysis in result.analysis_results:
            if analysis.summary:
                # 去重：同一内容的 summary 只写入一次
                if analysis.summary not in summaries:
                    summaries.append(analysis.summary)

        if not summaries:
            return f"本次分析处理了 {result.total_lines:,} 行日志数据，识别了 {result.statistics.get('by_level', {}).get('ERROR', 0):,} 个错误。"

        return summaries[0]  # 只取第一个 summary，不再拼接避免重复

    def _generate_combined_summary(self, results: List[ProcessingResult]) -> str:
        """生成多个文件的合并摘要"""
        file_summaries = []
        for result in results:
            summary = self._generate_summary(result)
            if summary:
                file_summaries.append(f"- **{os.path.basename(result.file_path)}**: {summary}")

        if not file_summaries:
            total_errors = sum(r.statistics.get('by_level', {}).get('ERROR', 0) for r in results)
            total_lines = sum(r.total_lines for r in results)
            return f"本次分析处理了 {len(results)} 个文件，共 {total_lines:,} 行日志，识别了 {total_errors:,} 个错误。"

        return "\n".join(file_summaries)

    def to_html(self, report: Report) -> str:
        """Generate HTML report with Apple-style design."""
        sections_html = ""
        
        for section in report.sections:
            if section.section_type == "suggestions":
                sections_html += self._generate_suggestions_html(section)
            elif section.section_type == "error_analysis":
                sections_html += self._generate_error_analysis_html(section)
            elif section.section_type == "statistics":
                sections_html += self._generate_statistics_html(section)
            else:
                sections_html += self._generate_default_section_html(section)
        
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{report.title}</title>
    <style>
        :root {{
            --primary: #007AFF;
            --primary-hover: #0066CC;
            --success: #34C759;
            --warning: #FF9500;
            --danger: #FF3B30;
            --bg: #FFFFFF;
            --bg-secondary: #F5F5F7;
            --bg-card: #FFFFFF;
            --bg-hover: #F8F8FA;
            --text: #1D1D1F;
            --text-secondary: #6E6E73;
            --text-tertiary: #8E8E93;
            --border: #E5E5EA;
            --border-light: #F0F0F0;
            --shadow: 0 1px 2px rgba(0, 0, 0, 0.05);
            --shadow-md: 0 4px 16px rgba(0, 0, 0, 0.08);
            --shadow-lg: 0 8px 32px rgba(0, 0, 0, 0.1);
            --radius: 16px;
            --radius-sm: 8px;
            --radius-md: 12px;
            --font: -apple-system, BlinkMacSystemFont, 'SF Pro Display', 'Helvetica Neue', sans-serif;
            --transition-fast: 0.15s ease;
            --transition-normal: 0.25s ease;
            --transition-slow: 0.35s ease;
        }}
        
        * {{ 
            margin: 0; 
            padding: 0; 
            box-sizing: border-box; 
        }}
        
        body {{ 
            font-family: var(--font); 
            background: var(--bg-secondary); 
            color: var(--text); 
            line-height: 1.5; 
            -webkit-font-smoothing: antialiased;
            -moz-osx-font-smoothing: grayscale;
        }}
        
        .header {{ 
            background: linear-gradient(180deg, #FFFFFF 0%, #F5F5F7 100%);
            border-bottom: 1px solid var(--border-light);
            padding: 32px 24px; 
            text-align: center;
        }}
        
        .header-content {{
            max-width: 800px;
            margin: 0 auto;
        }}
        
        .header h1 {{ 
            font-size: 28px; 
            font-weight: 600; 
            color: var(--text);
            margin-bottom: 8px;
            letter-spacing: -0.02em;
        }}
        
        .header p {{ 
            font-size: 14px; 
            color: var(--text-secondary);
            letter-spacing: 0.01em;
        }}
        
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 16px;
            padding: 24px;
            max-width: 1024px;
            margin: -48px auto 32px;
        }}
        
        .stat-card {{
            background: var(--bg-card);
            padding: 20px;
            border-radius: var(--radius);
            box-shadow: var(--shadow-md);
            text-align: center;
            transition: transform var(--transition-normal), box-shadow var(--transition-normal);
            border: 1px solid var(--border-light);
        }}
        
        .stat-card:hover {{
            transform: translateY(-2px);
            box-shadow: var(--shadow-lg);
        }}
        
        .stat-value {{ 
            font-size: 32px; 
            font-weight: 700; 
            color: var(--text);
            letter-spacing: -0.02em;
        }}
        
        .stat-label {{ 
            font-size: 13px; 
            color: var(--text-secondary); 
            margin-top: 8px;
            letter-spacing: 0.02em;
        }}
        
        .stat-value.success {{ color: var(--success); }}
        .stat-value.warning {{ color: var(--warning); }}
        .stat-value.danger {{ color: var(--danger); }}
        .stat-value.primary {{ color: var(--primary); }}
        
        .main {{ 
            max-width: 900px; 
            margin: 0 auto; 
            padding: 0 24px 48px; 
        }}
        
        .section-card {{
            background: var(--bg-card);
            border-radius: var(--radius);
            box-shadow: var(--shadow);
            border: 1px solid var(--border-light);
            margin-bottom: 20px;
            overflow: hidden;
            transition: box-shadow var(--transition-normal);
        }}
        
        .section-card:hover {{
            box-shadow: var(--shadow-md);
        }}
        
        .section-header {{
            padding: 20px 24px;
            border-bottom: 1px solid var(--border-light);
            cursor: pointer;
            display: flex;
            justify-content: space-between;
            align-items: center;
            transition: background-color var(--transition-fast);
        }}
        
        .section-header:hover {{ 
            background: var(--bg-hover); 
        }}
        
        .section-title {{ 
            font-weight: 600; 
            font-size: 15px; 
            color: var(--text);
            letter-spacing: -0.01em;
        }}
        
        .section-toggle {{ 
            color: var(--text-tertiary); 
            font-size: 12px;
            transition: transform var(--transition-normal);
            opacity: 0.6;
        }}
        
        .section-header:hover .section-toggle {{
            opacity: 1;
        }}
        
        .section-toggle.expanded {{ 
            transform: rotate(180deg); 
        }}
        
        .section-body {{ 
            padding: 24px;
            animation: fadeIn 0.3s ease;
        }}
        
        @keyframes fadeIn {{
            from {{
                opacity: 0;
                transform: translateY(-8px);
            }}
            to {{
                opacity: 1;
                transform: translateY(0);
            }}
        }}
        
        /* Markdown 转换样式 */
        .summary-text h1, .summary-text h2, .summary-text h3, .summary-text h4, .summary-text h5 {{
            margin: 24px 0 12px;
            font-weight: 600;
            color: var(--text);
            letter-spacing: -0.01em;
        }}
        
        .summary-text h1 {{ 
            font-size: 20px; 
            padding-bottom: 8px;
            border-bottom: 1px solid var(--border-light);
        }}
        
        .summary-text h2 {{ 
            font-size: 18px; 
        }}
        
        .summary-text h3 {{ 
            font-size: 16px; 
            color: var(--text-secondary);
        }}
        
        .summary-text h4 {{ 
            font-size: 15px; 
            color: var(--text-secondary);
            font-weight: 500;
        }}
        
        .summary-text h5 {{ 
            font-size: 14px; 
            color: var(--text-tertiary);
            font-weight: 500;
        }}
        
        .summary-text ul, .summary-text ol {{
            margin: 16px 0;
            padding-left: 24px;
        }}
        
        .summary-text ul {{
            list-style: disc;
        }}
        
        .summary-text ol {{
            list-style: decimal;
        }}
        
        .summary-text li {{
            margin: 8px 0;
            line-height: 1.7;
            font-size: 14px;
            color: var(--text-secondary);
        }}
        
        .summary-text strong {{
            font-weight: 600;
            color: var(--text);
        }}
        
        .summary-text code {{
            background: var(--bg-secondary);
            padding: 4px 10px;
            border-radius: var(--radius-sm);
            font-family: 'SF Mono', 'Monaco', 'Inconsolata', monospace;
            font-size: 13px;
            color: var(--text);
            letter-spacing: 0.01em;
        }}
        
        .summary-text pre {{
            background: #1D1D1F;
            color: #FFFFFF;
            padding: 16px;
            border-radius: var(--radius-md);
            overflow-x: auto;
            margin: 16px 0;
        }}
        
        .summary-text pre code {{
            background: transparent;
            color: #A7A7AA;
            padding: 0;
            font-size: 12px;
            line-height: 1.6;
        }}
        
        .summary-text p {{
            margin: 12px 0;
            line-height: 1.7;
            font-size: 14px;
            color: var(--text-secondary);
        }}
        
        .data-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 13px;
            margin: 16px 0;
        }}
        
        .data-table th {{
            text-align: left;
            padding: 12px 16px;
            background: var(--bg-secondary);
            font-weight: 600;
            font-size: 11px;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: var(--text-secondary);
            border-bottom: 1px solid var(--border);
        }}
        
        .data-table td {{
            padding: 12px 16px;
            border-bottom: 1px solid var(--border-light);
            color: var(--text-secondary);
            transition: background-color var(--transition-fast);
        }}
        
        .data-table tr:hover td {{ 
            background: var(--bg-hover); 
        }}
        
        .data-table tr:last-child td {{
            border-bottom: none;
        }}
        
        .severity-badge {{
            display: inline-block;
            padding: 4px 12px;
            border-radius: 9999px;
            font-size: 11px;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }}
        
        .severity-critical {{ 
            background: rgba(255, 59, 48, 0.1); 
            color: var(--danger); 
        }}
        
        .severity-high {{ 
            background: rgba(255, 149, 0, 0.1); 
            color: var(--warning); 
        }}
        
        .severity-medium {{ 
            background: rgba(0, 122, 255, 0.1); 
            color: var(--primary); 
        }}
        
        .severity-low {{ 
            background: rgba(52, 199, 89, 0.1); 
            color: var(--success); 
        }}
        
        .suggestions-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
            gap: 16px;
        }}
        
        .suggestion-card {{
            background: var(--bg-secondary);
            padding: 20px;
            border-radius: var(--radius-md);
            border-left: 4px solid var(--primary);
            transition: transform var(--transition-fast), box-shadow var(--transition-fast);
        }}
        
        .suggestion-card:hover {{
            transform: translateX(4px);
            box-shadow: var(--shadow);
        }}
        
        .suggestion-header {{
            font-weight: 600;
            margin-bottom: 12px;
            color: var(--text);
            font-size: 14px;
            letter-spacing: -0.01em;
        }}
        
        .suggestion-list {{
            list-style: none;
        }}
        
        .suggestion-list li {{
            padding: 10px 0;
            font-size: 13px;
            color: var(--text-secondary);
            border-bottom: 1px solid var(--border-light);
            line-height: 1.6;
        }}
        
        .suggestion-list li:last-child {{ 
            border-bottom: none; 
        }}
        
        .trend-list {{
            list-style: none;
            counter-reset: trend;
        }}
        
        .trend-list li {{
            position: relative;
            padding: 16px 0 16px 48px;
            border-bottom: 1px solid var(--border-light);
            counter-increment: trend;
            transition: background-color var(--transition-fast);
        }}
        
        .trend-list li:hover {{
            background: var(--bg-hover);
        }}
        
        .trend-list li::before {{
            content: counter(trend);
            position: absolute;
            left: 0;
            width: 32px;
            height: 32px;
            background: rgba(0, 122, 255, 0.08);
            color: var(--primary);
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 13px;
            font-weight: 600;
        }}
        
        .trend-list li:last-child {{
            border-bottom: none;
        }}
        
        .summary-text {{ 
            font-size: 14px; 
            line-height: 1.7; 
            color: var(--text-secondary); 
        }}
        
        .chart-bar-container {{
            display: flex;
            align-items: flex-end;
            gap: 8px;
            height: 120px;
            padding: 16px 0;
        }}
        
        .chart-bar {{
            flex: 1;
            background: linear-gradient(180deg, var(--primary) 0%, rgba(0, 122, 255, 0.3) 100%);
            border-radius: 6px 6px 0 0;
            position: relative;
            transition: opacity var(--transition-fast);
        }}
        
        .chart-bar:hover {{
            opacity: 0.8;
        }}
        
        .chart-bar-label {{
            position: absolute;
            bottom: -24px;
            left: 50%;
            transform: translateX(-50%);
            font-size: 11px;
            color: var(--text-tertiary);
            white-space: nowrap;
            letter-spacing: 0.02em;
        }}
        
        .highlight-box {{
            background: linear-gradient(135deg, rgba(0, 122, 255, 0.05) 0%, rgba(52, 199, 89, 0.05) 100%);
            border: 1px solid rgba(0, 122, 255, 0.1);
            border-radius: var(--radius-md);
            padding: 16px 20px;
            margin: 16px 0;
        }}
        
        .footer {{
            text-align: center;
            padding: 32px 24px;
            color: var(--text-tertiary);
            font-size: 12px;
            letter-spacing: 0.02em;
        }}
        
        @media (max-width: 768px) {{
            .stats-grid {{ 
                grid-template-columns: repeat(2, 1fr); 
                gap: 12px;
                padding: 16px;
                margin-top: -40px;
            }}
            
            .stat-card {{
                padding: 16px;
            }}
            
            .stat-value {{
                font-size: 24px;
            }}
            
            .suggestions-grid {{ 
                grid-template-columns: 1fr; 
            }}
            
            .section-body {{ 
                padding: 16px; 
            }}
            
            .section-header {{
                padding: 16px;
            }}
            
            .header {{
                padding: 24px 16px;
            }}
            
            .header h1 {{
                font-size: 22px;
            }}
            
            .main {{
                padding: 0 16px 32px;
            }}
        }}
        
        @media (max-width: 480px) {{
            .stats-grid {{ 
                grid-template-columns: 1fr; 
            }}
            
            .header h1 {{
                font-size: 18px;
            }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="header-content">
            <h1>📊 {report.title}</h1>
            <p>{report.generated_at.strftime('%Y-%m-%d %H:%M:%S')} | {report.file_size}</p>
        </div>
    </div>
    
    <div class="stats-grid">
        <div class="stat-card">
            <div class="stat-value primary">{report.total_lines:,}</div>
            <div class="stat-label">总行数</div>
        </div>
        <div class="stat-card">
            <div class="stat-value danger">{report.total_errors:,}</div>
            <div class="stat-label">错误数</div>
        </div>
        <div class="stat-card">
            <div class="stat-value warning">{report.total_warnings:,}</div>
            <div class="stat-label">警告数</div>
        </div>
        <div class="stat-card">
            <div class="stat-value success">100%</div>
            <div class="stat-label">处理进度</div>
        </div>
    </div>
    
    <div class="main">
        {sections_html}
        
        <div class="section-card">
            <div class="section-header">
                <span class="section-title">📝 总体摘要</span>
                <span class="section-toggle">▼</span>
            </div>
            <div class="section-body summary-text">
                {report.summary}
            </div>
        </div>
    </div>
    
    <div class="footer">
        Log Analyzer Report Generated by Trae AI
    </div>
    
    <script>
        document.querySelectorAll('.section-header').forEach(header => {{
            header.addEventListener('click', () => {{
                const body = header.nextElementSibling;
                const toggle = header.querySelector('.section-toggle');
                
                if (body.style.display === 'none' || !body.style.display) {{
                    body.style.display = 'block';
                    body.style.animation = 'fadeIn 0.3s ease';
                    toggle.classList.add('expanded');
                }} else {{
                    body.style.display = 'none';
                    toggle.classList.remove('expanded');
                }}
            }});
        }});
        
        document.querySelectorAll('.stat-card').forEach(card => {{
            card.addEventListener('mouseenter', () => {{
                card.style.transform = 'translateY(-4px)';
            }});
            
            card.addEventListener('mouseleave', () => {{
                card.style.transform = 'translateY(0)';
            }});
        }});
    </script>
</body>
</html>"""

    def _generate_default_section_html(self, section: ReportSection) -> str:
        content = markdown_to_html(section.content)
        return f"""
<div class="section-card">
    <div class="section-header">
        <span class="section-title">📄 {section.title}</span>
        <span class="section-toggle">▼</span>
    </div>
    <div class="section-body" style="display: block;">
        <div class="summary-text">{content}</div>
    </div>
</div>"""

    def _generate_statistics_html(self, section: ReportSection) -> str:
        data = section.data or {}
        by_level = data.get('错误级别分布', {})
        error_types = data.get('错误类型统计', {})
        top_classes = data.get('高频错误类', {})
        
        level_rows = ''.join([f'<tr><td>{k}</td><td>{v:,}</td></tr>' for k, v in by_level.items()])
        type_rows = ''.join([f'<tr><td>{k}</td><td>{v:,}</td></tr>' for k, v in list(error_types.items())[:10]])
        class_rows = ''.join([f'<tr><td>{k}</td><td>{v:,}</td></tr>' for k, v in list(top_classes.items())[:10]])
        
        return f"""
<div class="section-card">
    <div class="section-header">
        <span class="section-title">📊 {section.title}</span>
        <span class="section-toggle">▼</span>
    </div>
    <div class="section-body" style="display: block;">
        <h4 style="margin: 0 0 1rem; font-weight: 600;">错误级别分布</h4>
        <div class="chart-bar-container">
            {' '.join([f'<div class="chart-bar" style="height:{(v/max(list(by_level.values()) + [1]))*100}%"><div class="chart-bar-label">{k}</div></div>' for k, v in by_level.items()])}
        </div>
        
        <h4 style="margin: 1.5rem 0 1rem; font-weight: 600;">错误类型统计</h4>
        <table class="data-table">
            <thead><tr><th>类型</th><th>数量</th></tr></thead>
            <tbody>{type_rows}</tbody>
        </table>
        
        <h4 style="margin: 1.5rem 0 1rem; font-weight: 600;">高频错误类</h4>
        <table class="data-table">
            <thead><tr><th>类名</th><th>数量</th></tr></thead>
            <tbody>{class_rows}</tbody>
        </table>
    </div>
</div>"""

    def _generate_error_analysis_html(self, section: ReportSection) -> str:
        errors = section.data.get('关键错误', []) if section.data else []
        rows = ''
        for error in errors[:8]:
            severity_class = 'severity-' + error.get('severity', 'medium').lower()
            rows += f"""
<tr>
    <td>{error.get('error_type', '')}</td>
    <td>{error.get('description', '')}</td>
    <td>{error.get('count', 0):,}</td>
    <td><span class="severity-badge {severity_class}">{error.get('severity', '')}</span></td>
</tr>"""
        
        return f"""
<div class="section-card">
    <div class="section-header">
        <span class="section-title">🔴 {section.title}</span>
        <span class="section-toggle">▼</span>
    </div>
    <div class="section-body" style="display: block;">
        <table class="data-table">
            <thead><tr><th>错误类型</th><th>描述</th><th>次数</th><th>严重程度</th></tr></thead>
            <tbody>{rows}</tbody>
        </table>
    </div>
</div>"""

    def _generate_suggestions_html(self, section: ReportSection) -> str:
        data = section.data or {}
        ops = data.get('运维建议', [])
        dev = data.get('开发建议', [])
        
        ops_items = ''.join([f'<li>{item.get("suggestion", "")}</li>' for item in ops])
        dev_items = ''.join([f'<li>{item.get("suggestion", "")}</li>' for item in dev])
        
        return f"""
<div class="section-card">
    <div class="section-header">
        <span class="section-title">💡 {section.title}</span>
        <span class="section-toggle">▼</span>
    </div>
    <div class="section-body" style="display: block;">
        <div class="suggestions-grid">
            <div class="suggestion-card">
                <div class="suggestion-header">🔧 运维建议</div>
                <ul class="suggestion-list">{ops_items}</ul>
            </div>
            <div class="suggestion-card">
                <div class="suggestion-header">👨💻 开发建议</div>
                <ul class="suggestion-list">{dev_items}</ul>
            </div>
        </div>
    </div>
</div>"""

    def save_report(
        self,
        report: Report,
        format: str = "all",
        prefix: str = "report"
    ) -> List[str]:
        saved_files = []
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # 直接使用 prefix，不再从 report.file_path 提取文件名，避免重复
        report_name = f"{prefix}_{timestamp}"

        formats = format.split('+') if '+' in format else [format]
        
        need_json = format == "all" or "json" in formats
        need_md = format == "all" or "markdown" in formats or "md" in formats
        need_html = format == "all" or "html" in formats
        need_pdf = format == "all" or "pdf" in formats
        need_word = format == "all" or "word" in formats or "docx" in formats

        if need_json:
            json_path = os.path.join(self.output_dir, f"{report_name}.json")
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(report.to_dict(), f, indent=2, ensure_ascii=False)
            saved_files.append(json_path)

        if need_md:
            md_path = os.path.join(self.output_dir, f"{report_name}.md")
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(report.to_markdown())
            saved_files.append(md_path)

        if need_html:
            html_path = os.path.join(self.output_dir, f"{report_name}.html")
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(self.to_html(report))
            saved_files.append(html_path)

        if need_pdf:
            try:
                pdf_path = os.path.join(self.output_dir, f"{report_name}.pdf")
                # 使用类方法 _save_as_pdf，它接收 Report 对象
                self._save_as_pdf(report, pdf_path)
                saved_files.append(pdf_path)
            except ImportError:
                logging.warning("PDF 生成跳过: reportlab 模块未安装")
            except Exception as e:
                logging.warning(f"PDF 生成失败: {e}")

        if need_word:
            try:
                word_path = os.path.join(self.output_dir, f"{report_name}.docx")
                self._save_as_word(report, word_path)
                saved_files.append(word_path)
            except ImportError:
                logging.warning("Word 生成跳过: python-docx 模块未安装")
            except Exception as e:
                logging.warning(f"Word 生成失败: {e}")

        return saved_files

    def _convert_markdown_to_pdf_paragraph(self, md_content: str) -> str:
        """将 Markdown 格式转换为 PDF 支持的 HTML-like 格式"""
        if not md_content:
            return ""
        
        import re
        
        content = md_content
        
        # 先处理代码块（先保存再处理，避免干扰其他转换）
        code_blocks = []
        def save_code_block(match):
            code_blocks.append(match.group(0))
            return f"__CODE_BLOCK_PLACEHOLDER_{len(code_blocks)}__"
        content = re.sub(r'```[\s\S]*?```', save_code_block, content)
        
        # 先处理特殊字符（避免后续转义问题）
        # 但保留代码块占位符中的内容
        # 找到所有占位符位置
        placeholder_pattern = r'__CODE_BLOCK_PLACEHOLDER_\d+__'
        placeholders = {}
        for match in re.finditer(placeholder_pattern, content):
            placeholders[match.group(0)] = match.start()
        
        # 处理代码块以外的内容中的特殊字符
        result_parts = []
        last_end = 0
        for match in re.finditer(placeholder_pattern, content):
            # 处理占位符之前的内容
            before = content[last_end:match.start()]
            before = self._escape_html_chars(before)
            result_parts.append(before)
            # 保留占位符
            result_parts.append(match.group(0))
            last_end = match.end()
        # 处理最后的内容
        after = content[last_end:]
        after = self._escape_html_chars(after)
        result_parts.append(after)
        
        content = ''.join(result_parts)
        
        # 处理内联代码（必须在标题之前处理）
        content = re.sub(r'`([^`]+)`', r'<font name="Courier" color="#CC0000">\1</font>', content)
        
        # 正确处理粗体标记：**text** → <b>text</b>
        content = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', content)
        content = re.sub(r'__(.+?)__', r'<b>\1</b>', content)
        
        # 处理斜体
        content = re.sub(r'\*(.+?)\*', r'<i>\1</i>', content)
        content = re.sub(r'_(.+?)_', r'<i>\1</i>', content)
        
        # 处理标题（使用多行模式）
        content = re.sub(r'^#### (.+)$', r'<br/><b><font size="11">\1</font></b>', content, flags=re.MULTILINE)
        content = re.sub(r'^### (.+)$', r'<br/><b><font size="12">\1</font></b>', content, flags=re.MULTILINE)
        content = re.sub(r'^## (.+)$', r'<br/><br/><b><font size="14">\1</font></b>', content, flags=re.MULTILINE)
        content = re.sub(r'^# (.+)$', r'<br/><br/><b><font size="16">\1</font></b>', content, flags=re.MULTILINE)
        
        # 处理有序列表
        content = re.sub(r'^(\d+)\.\s+(.+)$', r'<br/>\1. \2', content, flags=re.MULTILINE)
        
        # 处理无序列表
        content = re.sub(r'^[-*+]\s+(.+)$', r'<br/>• \1', content, flags=re.MULTILINE)
        
        # 处理链接（简化处理）
        content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'<u>\1</u>', content)
        
        # 处理水平线
        content = re.sub(r'^[-*=_]{3,}$', r'<br/>────────────────────────────────────────<br/>', content, flags=re.MULTILINE)
        
        # 恢复代码块
        for i, code_block in enumerate(code_blocks):
            # 移除代码块标记和语言标识
            code_content = re.sub(r'```(\w+)?\n?', '', code_block)
            code_content = code_content.strip()
            # 处理代码内容中的特殊字符（但保留换行）
            code_content = code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            code_lines = code_content.split('\n')
            code_content_html = '<br/>'.join(code_lines)
            content = content.replace(f"__CODE_BLOCK_PLACEHOLDER_{i}__", 
                                     f'<br/><br/><font name="Courier" color="#333333" backcolor="#F5F5F5" fontSize="8">{code_content_html}</font><br/>')
        
        # 处理换行
        # 先合并多个连续换行
        content = re.sub(r'\n{3,}', '\n\n', content)
        # 将单个换行转换为 <br/>
        content = re.sub(r'\n', '<br/>', content)
        
        # 清理多余的 <br/>
        content = re.sub(r'(<br/>){3,}', '<br/><br/>', content)
        
        return content
    
    def _escape_html_chars(self, text: str) -> str:
        """转义 HTML 特殊字符"""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        return text


    def _convert_markdown_table_to_pdf(self, table_content: str) -> str:
        """将Markdown表格转换为PDF支持的文本格式"""
        import re
        
        lines = table_content.strip().split('\n')
        if len(lines) < 2:
            return ""
        
        # 解析表格
        header = lines[0].strip('|').split('|')
        body = lines[2:]
        
        # 移除空白字符
        header = [cell.strip() for cell in header]
        body_rows = []
        for line in body:
            cells = line.strip('|').split('|')
            cells = [cell.strip() for cell in cells]
            body_rows.append(cells)
        
        # 计算最大列宽
        max_cols = max(len(header), max(len(row) for row in body_rows))
        col_widths = [0] * max_cols
        
        for i, cell in enumerate(header):
            col_widths[i] = max(col_widths[i], len(cell))
        for row in body_rows:
            for i, cell in enumerate(row):
                if i < max_cols:
                    col_widths[i] = max(col_widths[i], len(cell))
        
        # 生成表格文本
        result = '<br/>'
        separator = '+' + '+'.join('-' * (w + 2) for w in col_widths) + '+'
        
        result += separator + '<br/>'
        
        # 表头
        row_cells = []
        for i, cell in enumerate(header):
            padding = col_widths[i] - len(cell)
            row_cells.append(f' {cell}{" " * padding} ')
        result += '|' + '|'.join(row_cells) + '|<br/>'
        
        result += separator + '<br/>'
        
        # 表体
        for row in body_rows:
            row_cells = []
            for i in range(max_cols):
                cell = row[i] if i < len(row) else ''
                padding = col_widths[i] - len(cell)
                row_cells.append(f' {cell}{" " * padding} ')
            result += '|' + '|'.join(row_cells) + '|<br/>'
        
        result += separator + '<br/>'
        
        return result

    def _add_data_to_pdf(self, story, data: Dict[str, Any], body_style) -> None:
        """将数据转换为格式化的PDF表格或文本，避免直接输出JSON"""
        from reportlab.platypus import Spacer, Table, TableStyle, Paragraph
        from reportlab.lib import colors
        from reportlab.lib.units import cm

        if not data:
            return

        for key, value in data.items():
            if not value:
                continue

            # 特殊处理处置动作建议数据
            if key == '处置动作建议' and isinstance(value, dict):
                self._add_response_actions_to_pdf(story, value, body_style)
                continue
            
            # 特殊处理根因推断数据
            if key == '根因推断' and isinstance(value, dict):
                self._add_root_cause_to_pdf(story, value, body_style)
                continue
            
            # 特殊处理因果链数据
            if key == '因果链' and isinstance(value, dict):
                self._add_causal_chain_to_pdf(story, value, body_style)
                continue
            
            # 特殊处理证据链数据
            if key == '证据链' and isinstance(value, dict):
                self._add_evidence_chain_to_pdf(story, value, body_style)
                continue
            
            # 特殊处理整改建议数据
            if key == '整改建议' and isinstance(value, dict):
                self._add_remediation_to_pdf(story, value, body_style)
                continue
            
            # 特殊处理解决建议数据
            if key == '解决建议' and isinstance(value, dict):
                self._add_suggestions_to_pdf(story, value, body_style)
                continue
            
            # 特殊处理趋势识别数据
            if key == '趋势识别' and isinstance(value, list):
                self._add_trends_to_pdf(story, value, body_style)
                continue
            
            # 特殊处理时间线数据
            if key == '时间线' and isinstance(value, dict):
                self._add_timeline_to_pdf(story, value, body_style)
                continue
            
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph(f"<b>{key}:</b>", body_style))

            if isinstance(value, dict):
                self._add_dict_table_to_pdf(story, value, body_style)
            elif isinstance(value, list):
                if len(value) > 0 and isinstance(value[0], dict):
                    self._add_list_table_to_pdf(story, value, body_style)
                else:
                    for item in value:
                        story.append(Paragraph(f"• {item}", body_style))
            elif isinstance(value, (int, float)):
                story.append(Paragraph(str(value), body_style))
            else:
                story.append(Paragraph(str(value), body_style))

            story.append(Spacer(1, 0.2 * cm))

    def _add_trends_to_pdf(self, story, trends: list, body_style) -> None:
        """将趋势识别数据转换为格式化的PDF内容"""
        from reportlab.platypus import Spacer, Paragraph
        from reportlab.lib.units import cm

        story.append(Spacer(1, 0.2 * cm))
        story.append(Paragraph("<b>趋势识别:</b>", body_style))
        story.append(Paragraph("通过对日志数据的深入分析，识别出以下关键趋势和模式：", body_style))
        story.append(Spacer(1, 0.1 * cm))
        
        for idx, trend in enumerate(trends[:10], 1):
            story.append(Paragraph(f"{idx}. <b>{trend}</b>", body_style))
        
        story.append(Spacer(1, 0.2 * cm))

    def _add_timeline_to_pdf(self, story, timeline: dict, body_style) -> None:
        """将时间线数据转换为格式化的PDF表格"""
        from reportlab.platypus import Spacer, Table, TableStyle, Paragraph
        from reportlab.lib import colors
        from reportlab.lib.units import cm

        story.append(Spacer(1, 0.2 * cm))
        story.append(Paragraph("<b>故障时间线:</b>", body_style))
        
        description = timeline.get('description', '')
        total_duration = timeline.get('total_duration', '')
        key_events = timeline.get('key_events', [])
        
        if description:
            story.append(Paragraph(f"<b>事件概述:</b> {description}", body_style))
        
        if total_duration:
            story.append(Paragraph(f"<b>持续时长:</b> {total_duration}", body_style))
        
        if key_events:
            story.append(Spacer(1, 0.1 * cm))
            story.append(Paragraph("<b>关键事件序列:</b>", body_style))
            
            event_type_names = {
                'first_abnormal': '🔴 首次异常',
                'peak_error': '🔥 错误峰值',
                'recovery': '🟢 恢复',
                'fault_confirmed': '⚠️ 故障确认'
            }
            
            table_data = [['序号', '时间', '事件类型', '描述']]
            for idx, event in enumerate(key_events, 1):
                event_time = event.get('time', 'N/A')
                event_type = event.get('event_type', 'unknown')
                event_desc = event.get('description', '')
                display_type = event_type_names.get(event_type, event_type)
                table_data.append([str(idx), event_time, display_type, event_desc])
            
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            
            story.append(table)
        
        story.append(Spacer(1, 0.2 * cm))

    def _add_root_cause_to_pdf(self, story, data: Dict[str, Any], body_style) -> None:
        """专门处理根因推断数据，以更易读的格式显示"""
        from reportlab.platypus import Spacer, Paragraph
        from reportlab.lib.units import cm
        
        if not data:
            return
        
        story.append(Spacer(1, 0.2 * cm))
        
        # 直接原因
        direct_cause = data.get('direct_cause', '')
        if direct_cause:
            story.append(Paragraph(f"<b>直接原因：</b>{direct_cause}", body_style))
        
        # 根本原因
        fundamental_cause = data.get('fundamental_cause', '')
        if fundamental_cause:
            story.append(Spacer(1, 0.15 * cm))
            story.append(Paragraph(f"<b>根本原因：</b>{fundamental_cause}", body_style))
        
        # 置信度
        confidence = data.get('confidence', '')
        if confidence:
            story.append(Spacer(1, 0.15 * cm))
            story.append(Paragraph(f"<b>置信度：</b>{confidence}", body_style))
        
        # 推理过程
        reasoning = data.get('reasoning', '')
        if reasoning:
            story.append(Spacer(1, 0.15 * cm))
            story.append(Paragraph(f"<b>推理过程：</b>{reasoning}", body_style))
    
    def _add_causal_chain_to_pdf(self, story, data: Dict[str, Any], body_style) -> None:
        """专门处理因果链数据，以更易读的格式显示"""
        from reportlab.platypus import Spacer, Paragraph
        from reportlab.lib.units import cm
        
        if not data:
            return
        
        story.append(Spacer(1, 0.2 * cm))
        
        # 因果链描述
        chain_description = data.get('chain_description', '')
        if chain_description:
            story.append(Paragraph(f"<b>因果链描述：</b>{chain_description}", body_style))
        
        # 因果传播路径
        chain_steps = data.get('chain_steps', [])
        if chain_steps:
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("<b>🔗 因果传播路径：</b>", body_style))
            
            for step in chain_steps:
                if isinstance(step, dict):
                    step_num = step.get('step', '')
                    cause = step.get('cause', '')
                    effect = step.get('effect', '')
                    evidence = step.get('evidence', '')
                    timestamp = step.get('timestamp', '')
                    
                    story.append(Spacer(1, 0.15 * cm))
                    if step_num:
                        story.append(Paragraph(f"<b>步骤 {step_num}：</b>", body_style))
                    if cause:
                        story.append(Paragraph(f"   • 原因：{cause}", body_style))
                    if effect:
                        story.append(Paragraph(f"   • 影响：{effect}", body_style))
                    if evidence:
                        story.append(Paragraph(f"   • 证据：{evidence}", body_style))
                    if timestamp:
                        story.append(Paragraph(f"   • 时间：{timestamp}", body_style))
                else:
                    story.append(Paragraph(f"   • {step}", body_style))
    
    def _add_evidence_chain_to_pdf(self, story, data: Dict[str, Any], body_style) -> None:
        """专门处理证据链数据，以更易读的格式显示"""
        from reportlab.platypus import Spacer, Paragraph
        from reportlab.lib.units import cm
        
        if not data:
            return
        
        story.append(Spacer(1, 0.2 * cm))
        
        # 证据链描述
        description = data.get('description', '')
        if description:
            story.append(Paragraph(f"<b>证据链描述：</b>{description}", body_style))
        
        # 证据列表
        evidences = data.get('evidences', [])
        if evidences:
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("<b>📋 证据列表：</b>", body_style))
            
            for i, evidence in enumerate(evidences, 1):
                if isinstance(evidence, dict):
                    timestamp = evidence.get('timestamp', '')
                    evidence_type = evidence.get('evidence_type', '')
                    content = evidence.get('content', '')
                    relevance = evidence.get('relevance', '')
                    
                    story.append(Spacer(1, 0.15 * cm))
                    story.append(Paragraph(f"<b>{i}. {evidence_type or '证据'}</b>", body_style))
                    if timestamp:
                        story.append(Paragraph(f"   • 时间：{timestamp}", body_style))
                    if content:
                        # 处理content可能是字典的情况
                        if isinstance(content, dict):
                            content_str = "; ".join(f"{k}: {v}" for k, v in content.items())
                        else:
                            content_str = str(content)
                        story.append(Paragraph(f"   • 内容：{content_str}", body_style))
                    if relevance:
                        story.append(Paragraph(f"   • 相关性：{relevance}", body_style))
                else:
                    story.append(Paragraph(f"   • {evidence}", body_style))

    def _add_remediation_to_pdf(self, story, data: Dict[str, Any], body_style) -> None:
        """专门处理整改建议数据，以更易读的格式显示"""
        from reportlab.platypus import Spacer, Paragraph
        from reportlab.lib.units import cm
        
        if not data:
            return
        
        story.append(Spacer(1, 0.2 * cm))
        
        # 立即处置
        immediate = data.get('immediate', [])
        if immediate:
            story.append(Paragraph("<b>⚡ 立即处置（Immediate Actions）</b>", body_style))
            story.append(Paragraph("<b>目标</b>: 快速恢复服务，最小化业务影响 | <b>时间要求</b>: 1小时内可执行", body_style))
            self._add_action_items_to_pdf(story, immediate, body_style)
        
        # 根本原因修复
        root_cause_fix = data.get('root_cause_fix', [])
        if root_cause_fix:
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("<b>🔧 根本原因修复（Root Cause Fix）</b>", body_style))
            story.append(Paragraph("<b>目标</b>: 消除问题根源，防止复发 | <b>时间要求</b>: 短期修复", body_style))
            self._add_action_items_to_pdf(story, root_cause_fix, body_style)
        
        # 架构与监控改进
        architecture_monitoring = data.get('architecture_monitoring', [])
        if architecture_monitoring:
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("<b>🏗️ 架构与监控改进（Architecture & Monitoring）</b>", body_style))
            story.append(Paragraph("<b>目标</b>: 增强系统韧性，提升可观测性 | <b>时间要求</b>: 中期优化", body_style))
            self._add_action_items_to_pdf(story, architecture_monitoring, body_style)
    
    def _add_action_items_to_pdf(self, story, actions: list, body_style) -> None:
        """将动作项列表以更易读的格式添加到PDF"""
        from reportlab.platypus import Spacer, Paragraph
        from reportlab.lib.units import cm
        
        for i, action in enumerate(actions, 1):
            if isinstance(action, dict):
                action_text = action.get('action', '')
                target = action.get('target', '')
                expected_effect = action.get('expected_effect', '')
                effort = action.get('effort_estimate', '')
                
                story.append(Spacer(1, 0.2 * cm))
                story.append(Paragraph(f"<b>{i}. {action_text}</b>", body_style))
                
                if target:
                    story.append(Paragraph(f"   • 目标对象：{target}", body_style))
                if expected_effect:
                    story.append(Paragraph(f"   • 预期效果：{expected_effect}", body_style))
                if effort:
                    story.append(Paragraph(f"   • 工作量预估：{effort}", body_style))
            else:
                story.append(Paragraph(f"   • {action}", body_style))
    
    def _add_suggestions_to_pdf(self, story, data: Dict[str, Any], body_style) -> None:
        """专门处理解决建议数据，以更易读的格式显示"""
        from reportlab.platypus import Spacer, Paragraph
        from reportlab.lib.units import cm
        
        if not data:
            return
        
        story.append(Spacer(1, 0.2 * cm))
        
        # 运维建议
        ops_suggestions = data.get('运维建议', [])
        if ops_suggestions:
            story.append(Paragraph("<b>🔹 运维建议</b>", body_style))
            self._add_suggestion_items_to_pdf(story, ops_suggestions, body_style)
        
        # 开发建议
        dev_suggestions = data.get('开发建议', [])
        if dev_suggestions:
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("<b>🔹 开发建议</b>", body_style))
            self._add_suggestion_items_to_pdf(story, dev_suggestions, body_style)
    
    def _add_suggestion_items_to_pdf(self, story, suggestions: list, body_style) -> None:
        """将建议项列表以更易读的格式添加到PDF"""
        from reportlab.platypus import Spacer, Paragraph
        from reportlab.lib.units import cm
        
        for i, suggestion in enumerate(suggestions, 1):
            if isinstance(suggestion, dict):
                category = suggestion.get('category', '')
                suggestion_text = suggestion.get('suggestion', '')
                
                story.append(Spacer(1, 0.15 * cm))
                if category:
                    story.append(Paragraph(f"<b>{i}. [{category}]</b> {suggestion_text}", body_style))
                else:
                    story.append(Paragraph(f"<b>{i}.</b> {suggestion_text}", body_style))
            else:
                story.append(Paragraph(f"   • {suggestion}", body_style))

    def _add_response_actions_to_pdf(self, story, data: Dict[str, Any], body_style) -> None:
        """专门处理处置动作建议数据，以更易读的格式显示"""
        from reportlab.platypus import Spacer, Paragraph
        from reportlab.lib.units import cm
        
        if not data:
            return
        
        # 处理描述
        description = data.get('description', '')
        if description:
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph(f"<b>描述：</b>{description}", body_style))
        
        # 处理应急止血动作
        emergency_actions = data.get('emergency_actions', [])
        if emergency_actions:
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("<b>🚨 应急止血动作：</b>", body_style))
            self._add_action_list_to_pdf(story, emergency_actions, body_style)
        
        # 处理故障排查动作
        troubleshooting_actions = data.get('troubleshooting_actions', [])
        if troubleshooting_actions:
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("<b>🔍 故障排查动作：</b>", body_style))
            self._add_action_list_to_pdf(story, troubleshooting_actions, body_style)
        
        # 处理恢复动作
        recovery_actions = data.get('recovery_actions', [])
        if recovery_actions:
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("<b>🔧 恢复动作：</b>", body_style))
            self._add_action_list_to_pdf(story, recovery_actions, body_style)
    
    def _add_action_list_to_pdf(self, story, actions: list, body_style) -> None:
        """将动作列表以更易读的格式添加到PDF"""
        from reportlab.platypus import Spacer, Paragraph
        from reportlab.lib.units import cm
        
        for i, action in enumerate(actions, 1):
            if isinstance(action, dict):
                action_name = action.get('action_name', '未命名动作')
                timing = action.get('timing', '')
                steps = action.get('steps', '')
                expected_effect = action.get('expected_effect', '')
                notes = action.get('notes', '')
                
                story.append(Spacer(1, 0.2 * cm))
                story.append(Paragraph(f"<b>{i}. {action_name}</b>", body_style))
                
                if timing:
                    story.append(Paragraph(f"   - 时机：{timing}", body_style))
                
                if steps:
                    # 处理步骤，可能是字符串或列表
                    if isinstance(steps, list):
                        steps_text = " → ".join(str(s) for s in steps)
                    else:
                        steps_text = str(steps)
                    story.append(Paragraph(f"   - 步骤：{steps_text}", body_style))
                
                if expected_effect:
                    story.append(Paragraph(f"   - 预期效果：{expected_effect}", body_style))
                
                if notes:
                    story.append(Paragraph(f"   - 备注：{notes}", body_style))
            else:
                story.append(Paragraph(f"   • {action}", body_style))

    def _add_dict_table_to_pdf(self, story, data: Dict[str, Any], body_style) -> None:
        """将字典数据转换为PDF表格"""
        from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
        from reportlab.lib import colors
        from reportlab.lib.units import cm

        if not data:
            return

        simple_items = []
        for k, v in data.items():
            if v is None:
                continue
            if isinstance(v, (str, int, float, bool)):
                simple_items.append([Paragraph(str(k), body_style), Paragraph(str(v), body_style)])
            elif isinstance(v, list) and len(v) <= 3:
                simple_items.append([Paragraph(str(k), body_style), Paragraph(', '.join(str(x) for x in v), body_style)])

        if not simple_items:
            story.append(Paragraph('<i>（详细数据见上方文本描述）</i>', body_style))
            return

        table_data = [[Paragraph('字段', body_style), Paragraph('值', body_style)]] + simple_items
        table = Table(table_data, colWidths=[6 * cm, 10 * cm])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), body_style.fontName),
            ('FONTNAME', (0, 1), (-1, -1), body_style.fontName),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(table)

    def _add_list_table_to_pdf(self, story, data_list: List[Dict[str, Any]], body_style) -> None:
        """将字典列表转换为PDF表格"""
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib import colors
        from reportlab.lib.units import cm

        if not data_list:
            return

        all_keys = set()
        for item in data_list:
            if isinstance(item, dict):
                all_keys.update(item.keys())

        priority_keys = ['error_type', 'description', 'severity', 'count',
                         'category', 'suggestion', 'message', 'timestamp',
                         'action', 'priority', 'status']
        columns = [k for k in priority_keys if k in all_keys]
        remaining = [k for k in all_keys if k not in columns]
        columns.extend(remaining[:max(0, 6 - len(columns))])

        if not columns:
            return

        table_data = [[Paragraph(str(c).replace('_', ' ').title(), body_style) for c in columns]]
        for item in data_list:
            if not isinstance(item, dict):
                continue
            row = []
            for key in columns:
                val = item.get(key, '')
                if isinstance(val, list):
                    cell_text = ', '.join(str(x) for x in val[:3])
                    if len(val) > 3:
                        cell_text += f' 等{len(val)}项'
                elif isinstance(val, dict):
                    cell_text = str(val)[:50] + '...' if len(str(val)) > 50 else str(val)
                else:
                    cell_text = str(val)[:100] + '...' if len(str(val)) > 100 else str(val)
                row.append(Paragraph(cell_text, body_style))
            table_data.append(row)

        col_width = 16 * cm / len(columns)
        table = Table(table_data, colWidths=[col_width] * len(columns))
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), body_style.fontName),
            ('FONTNAME', (0, 1), (-1, -1), body_style.fontName),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        story.append(table)

    def _save_as_pdf(self, report: Report, output_path: str) -> None:
        """将报告保存为 PDF 格式"""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        )
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os

        # 注册中文字体
        try:
            font_path = None
            subfont_index = 0
            
            # macOS 系统字体路径（TTC 集合文件需要指定 subfontIndex）
            mac_fonts = [
                ('/System/Library/Fonts/PingFang.ttc', 0),
                ('/System/Library/Fonts/PingFang.ttc', 1),
                ('/System/Library/Fonts/STHeiti Light.ttc', 0),
                ('/System/Library/Fonts/Supplemental/Songti.ttc', 0),
                ('/Library/Fonts/Arial Unicode.ttf', 0),
            ]
            # 其他系统字体路径
            other_fonts = [
                '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
                '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            ]
            
            for fp, idx in mac_fonts:
                if os.path.exists(fp):
                    font_path = fp
                    subfont_index = idx
                    break
            
            if not font_path:
                for fp in other_fonts:
                    if os.path.exists(fp):
                        font_path = fp
                        subfont_index = 0
                        break
            
            if font_path:
                pdfmetrics.registerFont(TTFont('Chinese', font_path, subfontIndex=subfont_index))
                chinese_font = 'Chinese'
            else:
                chinese_font = 'Helvetica'
                logging.warning("未找到中文字体，PDF 中的中文可能无法正常显示")
        except Exception as e:
            chinese_font = 'Helvetica'
            logging.warning(f"注册中文字体失败: {e}")

        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=2*cm,
            rightMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
            title=report.title
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=20,
            textColor=colors.HexColor('#1D1D1F'),
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName=chinese_font
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#007AFF'),
            spaceBefore=15,
            spaceAfter=10,
            fontName=chinese_font
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            spaceAfter=8,
            fontName=chinese_font
        )
        code_style = ParagraphStyle(
            'CodeStyle',
            parent=styles['Code'],
            fontSize=8,
            leading=10,
            fontName=chinese_font
        )

        story = []
        story.append(Paragraph(report.title, title_style))
        story.append(Paragraph(
            f"生成时间: {report.generated_at.strftime('%Y-%m-%d %H:%M:%S')}",
            body_style
        ))
        story.append(Paragraph(f"文件: {report.file_path}", body_style))
        story.append(Paragraph(f"文件大小: {report.file_size}", body_style))
        story.append(Spacer(1, 0.5*cm))

        stats_data = [
            ['总行数', f"{report.total_lines:,}"],
            ['错误数', f"{report.total_errors:,}"],
            ['警告数', f"{report.total_warnings:,}"],
            ['处理进度', '100%']
        ]
        stats_table = Table(stats_data, colWidths=[4*cm, 6*cm])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F5F5F7')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#1D1D1F')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), chinese_font),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E5E5EA'))
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 0.8*cm))

        for section in report.sections:
            # 章节标题（与Word/MD保持一致，不添加图标）
            story.append(Paragraph(section.title, heading_style))
            
            # 转换 Markdown 格式
            content = self._convert_markdown_to_pdf_paragraph(section.content)
            story.append(Paragraph(content, body_style))
            
            # 数据详情部分（与Word/MD保持一致，使用更美观的格式）
            if section.data:
                # 不再直接输出原始JSON，而是转换为格式化表格或文本
                self._add_data_to_pdf(story, section.data, body_style)
            
            story.append(Spacer(1, 0.6*cm))

        # 总体摘要（与Word/MD保持一致，不添加图标）
        if report.summary:
            story.append(Paragraph("7. 总体摘要", heading_style))
            summary_content = self._convert_markdown_to_pdf_paragraph(report.summary)
            story.append(Paragraph(summary_content, body_style))

        doc.build(story)

    def _add_markdown_paragraph_to_word(self, doc, md_content):
        """将 Markdown 格式内容添加到 Word 文档"""
        if not md_content:
            return
        
        import re
        
        # 分割段落
        paragraphs = re.split(r'\n\n+', md_content)
        
        for para in paragraphs:
            if not para.strip():
                continue
            
            # 检查是否为代码块
            if para.startswith('```'):
                # 移除代码块标记
                code_content = re.sub(r'```(\w+)?\n?', '', para)
                code_para = doc.add_paragraph(code_content.strip())
                code_para.style = 'No Spacing'
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                doc.add_paragraph()
                continue
            
            # 检查是否为表格
            if '|' in para and re.match(r'^\|.*\|$', para.split('\n')[0]):
                self._add_markdown_table_to_word(doc, para)
                doc.add_paragraph()
                continue
            
            # 检查是否为标题
            title_match = re.match(r'^(#{1,6})\s+(.+)', para)
            if title_match:
                level = len(title_match.group(1))
                doc.add_heading(title_match.group(2), level=level)
                continue
            
            # 检查是否为列表项
            list_match = re.match(r'^(\d+)\.\s+(.+)', para)
            if list_match:
                # 有序列表
                current_para = doc.add_paragraph()
                current_para.style = 'List Number'
                self._add_formatted_run_to_word(current_para, list_match.group(2))
                continue
            
            if para.startswith(('- ', '* ', '+ ')):
                # 无序列表
                current_para = doc.add_paragraph()
                current_para.style = 'List Bullet'
                self._add_formatted_run_to_word(current_para, para[2:])
                continue
            
            # 检查是否为水平线
            if re.match(r'^[-*=_]{3,}$', para.strip()):
                doc.add_paragraph().add_run('─' * 50).font.size = Pt(1)
                doc.add_paragraph()
                continue
            
            # 普通段落（处理行内格式）
            current_para = doc.add_paragraph()
            self._add_formatted_run_to_word(current_para, para)
    
    def _add_formatted_run_to_word(self, para, text):
        """向Word段落添加带格式的文本"""
        import re
        
        # 使用正则表达式匹配所有格式元素
        pattern = r'(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`|\[([^\]]+)\]\([^)]+\))'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            
            # 粗体
            if (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__')):
                run = para.add_run(part[2:-2])
                run.bold = True
            # 斜体
            elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
                run = para.add_run(part[1:-1])
                run.italic = True
            # 行内代码
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            # 链接
            elif '[' in part and '](' in part and part.endswith(')'):
                link_match = re.match(r'\[([^\]]+)\]\([^)]+\)', part)
                if link_match:
                    run = para.add_run(link_match.group(1))
                    run.underline = True
            else:
                para.add_run(part)
    
    def _add_markdown_table_to_word(self, doc, table_content):
        """将Markdown表格转换为Word表格"""
        import re
        
        lines = table_content.strip().split('\n')
        if len(lines) < 2:
            return
        
        # 解析表格
        header = lines[0].strip('|').split('|')
        separator = lines[1]
        body = lines[2:]
        
        # 移除空白字符
        header = [cell.strip() for cell in header]
        body_rows = []
        for line in body:
            cells = line.strip('|').split('|')
            cells = [cell.strip() for cell in cells]
            body_rows.append(cells)
        
        # 创建表格
        num_rows = len(body_rows) + 1
        num_cols = len(header)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # 设置表头
        for i, cell_text in enumerate(header):
            table.cell(0, i).text = cell_text
            for run in table.cell(0, i).paragraphs[0].runs:
                run.bold = True
        
        # 设置表体
        for i, row in enumerate(body_rows):
            for j, cell_text in enumerate(row):
                if j < num_cols:
                    table.cell(i + 1, j).text = cell_text

    def _add_data_to_word(self, doc, data: Dict[str, Any]) -> None:
        """将数据转换为格式化的Word表格或文本，避免直接输出JSON"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not data:
            return

        # 根据数据内容选择合适的展示方式
        for key, value in data.items():
            if not value:
                continue

            # 添加小标题
            doc.add_paragraph().add_run(f"{key}:").bold = True

            if isinstance(value, dict):
                # 字典数据 -> 两列表格（键值对）
                self._add_dict_table_to_word(doc, value)
            elif isinstance(value, list):
                if len(value) > 0 and isinstance(value[0], dict):
                    # 字典列表 -> 多列表格
                    self._add_list_table_to_word(doc, value)
                else:
                    # 普通列表 -> 项目符号列表
                    for item in value:
                        p = doc.add_paragraph()
                        p.add_run('• ').bold = True
                        p.add_run(str(item))
            elif isinstance(value, (int, float)):
                # 数字 -> 直接显示
                p = doc.add_paragraph()
                p.add_run(str(value))
            else:
                # 其他 -> 直接显示为文本
                p = doc.add_paragraph()
                p.add_run(str(value))

            # 每个数据项之间空一行
            doc.add_paragraph()

    def _add_dict_table_to_word(self, doc, data: Dict[str, Any]) -> None:
        """将字典数据转换为Word表格"""
        from docx.shared import Pt

        if not data:
            return

        # 过滤掉空值和嵌套复杂结构
        simple_items = []
        for k, v in data.items():
            if v is None:
                continue
            if isinstance(v, (str, int, float, bool)):
                simple_items.append((k, str(v)))
            elif isinstance(v, list) and len(v) <= 3:
                simple_items.append((k, ', '.join(str(x) for x in v)))

        if not simple_items:
            # 如果没有简单项，显示提示
            p = doc.add_paragraph()
            p.add_run('（详细数据见上方文本描述）').italic = True
            p.runs[0].font.size = Pt(9)
            return

        table = doc.add_table(rows=len(simple_items), cols=2)
        table.style = 'Light Grid Accent 1'

        for i, (k, v) in enumerate(simple_items):
            table.cell(i, 0).text = str(k)
            table.cell(i, 1).text = str(v)
            # 设置字体
            for cell in (table.cell(i, 0), table.cell(i, 1)):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(10)
            # 第一列加粗
            for run in table.cell(i, 0).paragraphs[0].runs:
                run.bold = True

    def _add_list_table_to_word(self, doc, data_list: List[Dict[str, Any]]) -> None:
        """将字典列表转换为Word表格"""
        from docx.shared import Pt

        if not data_list:
            return

        # 收集所有可能的列名
        all_keys = set()
        for item in data_list:
            if isinstance(item, dict):
                all_keys.update(item.keys())

        # 优先保留常见字段，并限制列数
        priority_keys = ['error_type', 'description', 'severity', 'count',
                         'category', 'suggestion', 'message', 'timestamp',
                         'action', 'priority', 'status']
        columns = [k for k in priority_keys if k in all_keys]
        # 补充其他字段（最多6列）
        remaining = [k for k in all_keys if k not in columns]
        columns.extend(remaining[:max(0, 6 - len(columns))])

        if not columns:
            return

        table = doc.add_table(rows=len(data_list) + 1, cols=len(columns))
        table.style = 'Light Grid Accent 1'

        # 表头
        for j, col_name in enumerate(columns):
            cell = table.cell(0, j)
            cell.text = str(col_name).replace('_', ' ').title()
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(10)

        # 数据行
        for i, item in enumerate(data_list):
            if not isinstance(item, dict):
                continue
            for j, key in enumerate(columns):
                val = item.get(key, '')
                # 处理不同类型的值
                if isinstance(val, list):
                    cell_text = ', '.join(str(x) for x in val[:3])
                    if len(val) > 3:
                        cell_text += f' 等{len(val)}项'
                elif isinstance(val, dict):
                    cell_text = str(val)[:50] + '...' if len(str(val)) > 50 else str(val)
                else:
                    cell_text = str(val)[:100] + '...' if len(str(val)) > 100 else str(val)

                cell = table.cell(i + 1, j)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Microsoft YaHei'
                        run.font.size = Pt(9)

    def _save_as_word(self, report: Report, output_path: str) -> None:
        """将报告保存为 Word 格式"""
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        
        title = doc.add_heading(report.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_para = doc.add_paragraph()
        info_para.add_run(f"生成时间: {report.generated_at.strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
        info_para.add_run(f"文件: {report.file_path}\n").italic = True
        info_para.add_run(f"文件大小: {report.file_size}").italic = True

        doc.add_paragraph()
        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        stats_data = [
            ('总行数', f"{report.total_lines:,}"),
            ('错误数', f"{report.total_errors:,}"),
            ('警告数', f"{report.total_warnings:,}"),
            ('处理进度', '100%')
        ]
        for i, (k, v) in enumerate(stats_data):
            stats_table.cell(i, 0).text = k
            stats_table.cell(i, 1).text = v

        doc.add_paragraph()

        for section in report.sections:
            # 移除图标，与PDF格式保持一致
            doc.add_heading(section.title, level=1)
            # 使用新方法处理 Markdown 格式
            self._add_markdown_paragraph_to_word(doc, section.content)
            
            # 数据详情部分（与PDF格式保持一致）
            # 注意：不再直接输出原始JSON，而是将数据转换为格式化的表格或文本
            if section.data:
                self._add_data_to_word(doc, section.data)

        # 总体摘要（与PDF格式保持一致）
        if report.summary:
            doc.add_heading("7. 总体摘要", level=1)
            self._add_markdown_paragraph_to_word(doc, report.summary)

        doc.save(output_path)

    def _create_combined_overview_section(self, results: List[ProcessingResult]) -> ReportSection:
        """创建综合概览部分"""
        completed = sum(1 for r in results if r.status == 'completed')
        total_lines = sum(r.total_lines for r in results)
        total_errors = sum(
            r.statistics.get('by_level', {}).get('ERROR', 0) +
            r.statistics.get('by_level', {}).get('FATAL', 0)
            for r in results
        )
        
        content = f"""
**分析文件数**: {len(results)} 个
**已完成**: {completed}/{len(results)}
**总行数**: {total_lines:,}
**总错误数**: {total_errors:,}

### 文件列表
"""
        for i, result in enumerate(results, 1):
            content += f"{i}. `{os.path.basename(result.file_path)}` - {result.status}\n"

        return ReportSection(
            title="综合概览",
            content=content.strip(),
            section_type="overview",
            data={
                'total_files': len(results),
                'completed_files': completed,
                'total_lines': total_lines,
                'total_errors': total_errors
            }
        )

    def _create_file_summaries_section(self, results: List[ProcessingResult]) -> ReportSection:
        """创建各文件摘要部分"""
        content = ""
        for result in results:
            if result.status != 'completed':
                continue
                
            file_errors = (
                result.statistics.get('by_level', {}).get('ERROR', 0) +
                result.statistics.get('by_level', {}).get('FATAL', 0)
            )
            file_warnings = result.statistics.get('by_level', {}).get('WARN', 0)
            
            content += f"### {os.path.basename(result.file_path)}\n"
            content += f"- 状态: {result.status}\n"
            content += f"- 行数: {result.total_lines:,}\n"
            content += f"- 错误数: {file_errors:,}\n"
            content += f"- 警告数: {file_warnings:,}\n"
            
            # 添加分析摘要
            summaries = [a.summary for a in result.analysis_results if a.summary]
            if summaries:
                content += f"- 摘要: {summaries[0][:100]}...\n"
            
            content += "\n"

        return ReportSection(
            title="各文件分析摘要",
            content=content.strip(),
            section_type="summaries"
        )

    def _create_combined_statistics_section(self, results: List[ProcessingResult]) -> ReportSection:
        """创建综合统计部分"""
        all_stats = {
            'by_level': {'ERROR': 0, 'WARN': 0, 'INFO': 0, 'DEBUG': 0, 'FATAL': 0},
            'error_types': {},
            'top_classes': {}
        }

        for result in results:
            stats = result.statistics
            for level, count in stats.get('by_level', {}).items():
                if level in all_stats['by_level']:
                    all_stats['by_level'][level] += count
            
            for error_type, count in stats.get('error_types', {}).items():
                all_stats['error_types'][error_type] = all_stats['error_types'].get(error_type, 0) + count
            
            for class_name, count in stats.get('top_classes', {}).items():
                all_stats['top_classes'][class_name] = all_stats['top_classes'].get(class_name, 0) + count

        content = "### 错误级别分布\n\n"
        for level in ['ERROR', 'FATAL', 'WARN', 'INFO', 'DEBUG']:
            content += f"- {level}: {all_stats['by_level'][level]:,}\n"

        content += "\n### 错误类型Top 10\n\n"
        top_errors = sorted(all_stats['error_types'].items(), key=lambda x: x[1], reverse=True)[:10]
        for error_type, count in top_errors:
            content += f"- {error_type}: {count:,}\n"

        content += "\n### 涉及类Top 10\n\n"
        top_classes = sorted(all_stats['top_classes'].items(), key=lambda x: x[1], reverse=True)[:10]
        for class_name, count in top_classes:
            content += f"- {class_name}: {count:,}\n"

        return ReportSection(
            title="综合统计分析",
            content=content.strip(),
            section_type="statistics",
            data=all_stats
        )

    def _create_combined_suggestions_section(self, results: List[ProcessingResult]) -> ReportSection:
        """创建综合建议部分"""
        all_suggestions = []
        
        for result in results:
            for analysis in result.analysis_results:
                if analysis.suggestions:
                    all_suggestions.extend(analysis.suggestions)
        
        # 去重并排序
        unique_suggestions = list(dict.fromkeys(all_suggestions))
        
        content = "基于所有文件的分析，以下是综合解决建议：\n\n"
        for idx, suggestion in enumerate(unique_suggestions[:5], 1):
            content += f"{idx}. {suggestion}\n"

        return ReportSection(
            title="综合解决建议",
            content=content.strip(),
            section_type="suggestions",
            data={'total_suggestions': len(unique_suggestions)}
        )

    def _generate_combined_summary(self, results: List[ProcessingResult]) -> str:
        """生成综合摘要"""
        completed = sum(1 for r in results if r.status == 'completed')
        total_errors = sum(
            r.statistics.get('by_level', {}).get('ERROR', 0) +
            r.statistics.get('by_level', {}).get('FATAL', 0)
            for r in results
        )
        
        summary = f"本次综合分析处理了 {len(results)} 个文件（{completed} 个成功），"
        summary += f"共 {sum(r.total_lines for r in results):,} 行日志，"
        summary += f"识别了 {total_errors:,} 个错误。"
        
        return summary

    def generate_batch_reports(
        self,
        results: List[ProcessingResult],
        format: str = "both"
    ) -> Dict[str, List[str]]:
        all_reports = {}

        for result in results:
            if result.status == "completed":
                report = self.generate_report(result)
                saved_files = self.save_report(report, format)
                all_reports[result.file_path] = saved_files

        return all_reports
